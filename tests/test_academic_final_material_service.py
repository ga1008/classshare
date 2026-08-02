from __future__ import annotations

import asyncio
import json
import sqlite3
import unittest
from datetime import datetime, timedelta
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock, patch
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
import httpx
from PIL import Image

from classroom_app.db import schema_academic_final_materials
from classroom_app.services.academic_exam_roster_sync_service import _exam_course_from_row
from classroom_app.services.academic_final_material_document_service import (
    build_exam_analysis_docx,
    build_grade_register_docx,
)
from classroom_app.services.academic_final_material_service import (
    ACADEMIC_EXAM_ANALYSIS_TYPE,
    ACADEMIC_GRADE_REGISTER_TYPE,
    academic_final_material_record_urls,
    build_exam_analysis_export_payload,
    build_grade_register_export_payload,
    is_grade_entry_submitted,
    list_teacher_final_material_candidates,
    _download_fine_report_word,
    parse_exam_analysis_rtf,
    parse_grade_register_rtf,
    repair_legacy_grade_register_roster_order,
    reclaim_stale_academic_final_material_batches,
    upsert_batch_state,
    validate_paired_reports,
)
from classroom_app.services.manage_nav_service import build_manage_nav
from classroom_app.services.material_ai_import_service import resolve_material_ai_import_type
from classroom_app.services.material_export_template_service import build_material_export_artifact


def _rtf_bytes(text: str) -> bytes:
    parts = [r"{\rtf1\ansi\uc1 "]
    for char in text:
        if char == "\n":
            parts.append(r"\par ")
        elif char == "\t":
            parts.append(r"\tab ")
        elif char in r"\{}":
            parts.append("\\" + char)
        elif ord(char) < 128:
            parts.append(char)
        else:
            codepoint = ord(char)
            signed = codepoint if codepoint < 32768 else codepoint - 65536
            parts.append(rf"\u{signed}?")
    parts.append("}")
    return "".join(parts).encode("ascii")


def _gb18030_hex_rtf(text: str) -> bytes:
    parts = [r"{\rtf1\ansi\ansicpg936 "]
    for value in text.encode("gb18030"):
        if value == 9:
            parts.append(r"\tab ")
        elif value == 10:
            parts.append(r"\par ")
        elif value in b"\\{}":
            parts.append("\\" + chr(value))
        elif 32 <= value < 127:
            parts.append(chr(value))
        else:
            parts.append(rf"\'{value:02x}")
    parts.append("}")
    return "".join(parts).encode("ascii")


GRADE_TEXT = """广西外国语学院期末成绩登记表
2025-2026学年第2学期
开课部门\t软件工程系\t班级\t软工2303班\t任课教师\t张老师\t学分\t2.0
课程名称\t计算机网络实验\t课程性质\t实践教学\t考核方式\t考查\t填表日期\t2026-07-31
学号\t姓名\t平时\t期中\t实验在线\t期末\t总评\t备注\t学号\t姓名\t平时\t期中\t实验在线\t期末\t总评\t备注
2300000001\t学生甲\t100\t\t\t80\t88\t\t2300000002\t学生乙\t100\t\t\t100\t100\t
教师：
总评成绩 = 平时*40% + 期末*60%"""


ANALYSIS_TEXT = """广西外国语学院课程试卷分析表
2025-2026学年第2学期
课程名称\t计算机网络实验\t学时数\t32\t开课单位\t软件工程系
教师姓名\t张老师
学生班级\t软工2303班
人数\t0\t0\t0\t1\t1
比例\t0.00\t0.00\t0.00\t50.00\t50.00
平均分\t90.00\t标准差\t14.14
最高分\t100\t最低分\t80\t及格率\t100.00"""


class AcademicFinalMaterialServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.grade = parse_grade_register_rtf(_rtf_bytes(GRADE_TEXT))
        self.analysis = parse_exam_analysis_rtf(_rtf_bytes(ANALYSIS_TEXT))
        self.validation = validate_paired_reports(
            self.grade,
            self.analysis,
            remote_student_count=2,
        )

    def test_generated_document_urls_are_release_and_record_versioned(self) -> None:
        with patch(
            "classroom_app.services.academic_final_material_service.get_deployment_release_id",
            return_value="release-42",
        ):
            first = academic_final_material_record_urls(36, "2026-08-03T01:00:00")
            second = academic_final_material_record_urls(36, "2026-08-03T02:00:00")
        self.assertIn("format=docx&v=release-42-", first["export_url"])
        self.assertIn("format=docx&v=release-42-", first["preview_url"])
        self.assertNotEqual(first["export_url"], second["export_url"])

    def test_paired_parser_and_deterministic_validator_pass(self) -> None:
        self.assertEqual(2, len(self.grade["students"]))
        self.assertEqual("计算机网络实验", self.grade["fields"]["course_name"])
        self.assertTrue(self.validation["passed"], self.validation["errors"])
        self.assertEqual([0, 0, 0, 1, 1], self.validation["computed"]["distribution_counts"])
        self.assertEqual(14.14, self.validation["computed"]["standard_deviation"])

    def test_validator_blocks_modified_analysis_statistics(self) -> None:
        self.analysis["statistics"]["average"] = 99.99
        validation = validate_paired_reports(self.grade, self.analysis, remote_student_count=2)
        self.assertFalse(validation["passed"])
        check = next(item for item in validation["checks"] if item["key"] == "statistics_average")
        self.assertFalse(check["ok"])
        self.assertIn("平均分", validation["errors"][0])

    def test_validator_treats_course_name_case_and_width_as_equivalent(self) -> None:
        context = {"course_name": "动态ｗｅｂ程序设计"}
        self.grade["fields"]["course_name"] = "动态Web程序设计"
        self.analysis["fields"]["course_name"] = "动态 WEB 程序设计"

        validation = validate_paired_reports(
            self.grade,
            self.analysis,
            context=context,
            remote_student_count=2,
        )

        self.assertTrue(validation["passed"], validation["errors"])
        context_check = next(item for item in validation["checks"] if item["key"] == "context_course_name")
        self.assertTrue(context_check["ok"])

    def test_grade_entry_status_does_not_treat_negative_labels_as_submitted(self) -> None:
        for status in ("未提交", "未录入", "待提交", "未完成"):
            self.assertFalse(is_grade_entry_submitted(status), status)
        for status in ("已提交", "已录入", "已完成"):
            self.assertTrue(is_grade_entry_submitted(status), status)
        self.assertTrue(is_grade_entry_submitted("", {"cjsftj": "1"}))
        self.assertTrue(
            is_grade_entry_submitted(
                "3",
                {
                    "lrzt": "3",
                    "lrztmc": "提交",
                    "cjlrshzt": "审核通过",
                    "tjsj": "2026-07-31 02:08:29",
                },
            )
        )

    def test_exam_course_prefers_human_readable_submission_status(self) -> None:
        course = _exam_course_from_row(
            {
                "jxb_id": "teaching-class-1",
                "kch_id": "course-1",
                "kcmc": "计算机网络实验",
                "lrzt": "3",
                "lrztmc": "提交",
                "cjlrshzt": "审核通过",
            },
            source_url="https://jwxt.gxufl.com/example",
            term_params={"xnm": "2025", "xqm": "12"},
        )
        self.assertIsNotNone(course)
        self.assertEqual("提交", course.grade_entry_status)

    def test_both_docx_renderers_produce_openable_documents(self) -> None:
        grade_payload = build_grade_register_export_payload(self.grade, self.validation)
        analysis_payload = build_exam_analysis_export_payload(
            self.analysis,
            self.validation,
            defaults={
                "proposition_form": "教师组题",
                "exam_form": "闭卷",
                "separate_teaching_exam": "否",
                "course_nature": "必修",
                "marking_form": "本人阅卷",
                "analysis_text": "试题结构合理，覆盖核心知识与实践能力。成绩分布集中且整体掌握良好，后续将通过分层案例、课堂复盘和专项训练强化综合分析与知识迁移。",
            },
        )
        for payload, builder, title in (
            (grade_payload, build_grade_register_docx, "广西外国语学院期末成绩登记表"),
            (analysis_payload, build_exam_analysis_docx, "广西外国语学院课程试卷分析表"),
        ):
            content = builder(payload)
            self.assertTrue(content.startswith(b"PK"))
            document = Document(BytesIO(content))
            all_text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            self.assertIn(title, all_text)
            self.assertIn("计算机网络实验", all_text)

    def test_exam_analysis_rebuilds_official_single_table_geometry(self) -> None:
        payload = build_exam_analysis_export_payload(
            self.analysis,
            self.validation,
            defaults={
                "proposition_form": "教师组题",
                "exam_form": "闭卷",
                "separate_teaching_exam": "否",
                "course_nature": "必修",
                "marking_form": "本人阅卷",
                "analysis_text": "一、成绩分布与试卷分析\n本次成绩能够反映课程目标达成情况。\n二、改进意见与措施\n后续加强综合实践训练。",
            },
        )
        with TemporaryDirectory() as temp_dir:
            department_signature = Path(temp_dir) / "department.png"
            dean_signature = Path(temp_dir) / "dean.png"
            Image.new("RGBA", (220, 72), (0, 0, 0, 0)).save(department_signature)
            Image.new("RGBA", (220, 72), (1, 1, 1, 0)).save(dean_signature)
            payload["fields"]["department_signature_image_path"] = str(department_signature)
            payload["fields"]["dean_signature_image_path"] = str(dean_signature)
            content = build_exam_analysis_docx(payload)

        document = Document(BytesIO(content))
        self.assertEqual(1, len(document.tables))
        self.assertEqual(22, len(document.tables[0].rows))
        section = document.sections[0]
        self.assertEqual(11905, section.page_width.twips)
        self.assertEqual(16837, section.page_height.twips)
        self.assertEqual(388, section.top_margin.twips)
        self.assertEqual(388, section.bottom_margin.twips)
        self.assertEqual(1080, section.left_margin.twips)
        self.assertEqual(1080, section.right_margin.twips)

        table = document.tables[0]
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        self.assertEqual([434, 930, 630, 1080, 570, 870, 1425, 915, 1080, 840, 870], grid_widths)
        row_heights = [
            int(row._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")))
            if row._tr.get_or_add_trPr().find(qn("w:trHeight")) is not None else None
            for row in table.rows
        ]
        self.assertEqual([720, 340, 220], row_heights[:3])
        self.assertEqual(3000, row_heights[15])
        self.assertEqual(4960, row_heights[17])
        self.assertEqual(1542, row_heights[20])
        self.assertEqual("广西外国语学院课程试卷分析表", table.rows[0].cells[0].text)
        self.assertIn("教师组题", table.rows[5].cells[7].text)
        self.assertEqual("√", table.rows[5].cells[9].text)
        self.assertIn("本人阅卷 √", table.rows[13].cells[2].text)
        self.assertIn("简要分析试题结构", table.rows[16].cells[1].text)
        self.assertIn("成绩分布与试卷分析", table.rows[17].cells[1].text)
        self.assertIn("本表一式两份", table.rows[21].cells[0].text)
        self.assertIn('w:textDirection w:val="tbRl"', document._element.xml)

        with ZipFile(BytesIO(content)) as archive:
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertGreaterEqual(len(media), 2)  # Word may deduplicate identical signature images.
        self.assertEqual(3, document_xml.count("<a:blip "))
        self.assertNotIn(str(department_signature), document_xml)
        self.assertNotIn(str(dean_signature), document_xml)
        self.assertNotIn("已阅", document_xml)
        self.assertNotIn("{{", document_xml)

    def test_exam_analysis_parser_repairs_gb18030_ansi_and_reads_checkmarks(self) -> None:
        text = """广西外国语学院课程试卷分析表
2025-2026学年第1学期
课程名称\t服务器配置与管理\t学时数\t48\t开课单位\tE02软件工程系
教师姓名\t张海林\t课程性质\t选修\t√\t必修
命题形式(打√)\t试题库\t\t试卷库\t\t教师组题\t√
考试形式(打√)\t开卷\t√\t闭卷\t\t教考分离(打√)\t是\t\t否\t√
学生班级\t软工2406班（专升本）
人数\t0\t2\t4\t5\t34
比例\t0.00%\t4.44%\t8.89%\t11.11%\t75.56%
平均分\t88.71\t标准差\t7.75
最高分\t98\t最低分\t63\t及格率\t100.00%
阅卷形式(打√)\t本人阅卷 √\t同行阅卷\t集体阅卷\t机器阅卷\t其他"""
        body = text.encode("gb18030").replace(b"\t", b"\\tab ").replace(b"\n", b"\\par ")
        parsed = parse_exam_analysis_rtf(b"{\\rtf1\\ansi\\ansicpg936 " + body + b"}")
        self.assertEqual("服务器配置与管理", parsed["fields"]["course_name"])
        self.assertEqual("选修", parsed["fields"]["course_nature"])
        self.assertEqual("教师组题", parsed["fields"]["proposition_form"])
        self.assertEqual("开卷", parsed["fields"]["exam_form"])
        self.assertEqual("否", parsed["fields"]["separate_teaching_exam"])
        self.assertEqual("本人阅卷", parsed["fields"]["marking_form"])

    def test_grade_register_reuses_official_single_table_geometry_and_slot_order(self) -> None:
        payload = build_grade_register_export_payload(self.grade, self.validation)
        students = []
        for index in range(45):
            score = float(100 - index)
            students.append(
                {
                    "student_number": f"24000000{index + 1:03d}",
                    "student_name": f"学生{index + 1:02d}",
                    "ordinary_score": score,
                    "midterm_score": None,
                    "experiment_online_score": None,
                    "final_exam_score": score,
                    "final_score": score,
                    "remark": "",
                }
            )
        payload["structured"]["students"] = students
        payload["structured"]["statistics"] = {"student_count": 45, "average": 78.0}

        content = build_grade_register_docx(payload)
        document = Document(BytesIO(content))
        self.assertEqual(1, len(document.sections))
        self.assertEqual(1, len(document.tables))
        section = document.sections[0]
        self.assertEqual(11905, section.page_width.twips)
        self.assertEqual(16837, section.page_height.twips)
        self.assertEqual(226, section.top_margin.twips)
        self.assertEqual(5, section.bottom_margin.twips)
        self.assertEqual(283, section.left_margin.twips)
        self.assertEqual(283, section.right_margin.twips)

        table = document.tables[0]
        self.assertEqual(47, len(table.rows))
        self.assertEqual(17, len(table.columns))
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        self.assertEqual(
            [1473, 850, 566, 510, 566, 567, 567, 511, 396, 1078, 851, 567, 511, 567, 567, 567, 511],
            grid_widths,
        )
        row_heights = [
            int(row._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")))
            for row in table.rows
        ]
        self.assertEqual([480, 280, 340, 340, 500], row_heights[:5])
        self.assertEqual([180, 280], row_heights[-2:])
        self.assertEqual("学生40", table.rows[44].cells[1].text)
        self.assertEqual("学生41", table.rows[5].cells[10].text)
        self.assertEqual("学生45", table.rows[9].cells[10].text)
        self.assertEqual("", table.rows[10].cells[10].text)
        self.assertIn("总评成绩 =", table.rows[35].cells[9].text)
        self.assertEqual("期末成绩分析表", table.rows[36].cells[8].text)
        self.assertEqual(
            [11, 10, 10, 10, 4, 0, 0, 0],
            [int(table.rows[row].cells[10].text) for row in range(37, 45)],
        )
        self.assertNotIn("F3F4F6", document._element.xml)
        self.assertNotIn("{{", document._element.xml)

        with ZipFile(BytesIO(content)) as archive:
            self.assertNotIn("word/media/image1.png", archive.namelist())
            document_xml = archive.read("word/document.xml").decode("utf-8")
        for reference_value in ("甘鸿明", "张海林", "服务器配置与管理", "软工2406"):
            self.assertNotIn(reference_value, document_xml)

    def test_grade_parser_repairs_ansi_hex_and_preserves_column_major_roster_order(self) -> None:
        source = """广西外国语学院期末成绩登记表
2025-2026学年第1学期
开课部门：E02软件工程系\t班级：软工2406班\t任课教师：张海林\t学分：3.0
课程名称：服务器配置与管理\t课程性质：专业教育\t考核方式：考试\t填表日期：2026-03-05
学号\t姓名\t平时\t期中\t实验/在线\t期末\t总评\t备注\t学号\t姓名\t平时\t期中\t实验/在线\t期末\t总评\t备注
24000000001\t左一\t90\t\t\t90\t90\t\t24000000003\t右一\t80\t\t\t80\t80\t
24000000002\t左二\t91\t\t\t91\t91
教师："""
        parsed = parse_grade_register_rtf(_gb18030_hex_rtf(source))
        self.assertEqual("服务器配置与管理", parsed["fields"]["course_name"])
        self.assertEqual(
            ["24000000001", "24000000002", "24000000003"],
            [item["student_number"] for item in parsed["students"]],
        )
        self.assertEqual({"left_student_count": 2, "right_student_count": 1}, parsed["source_layout"])

    def test_legacy_persisted_roster_is_lazily_repaired_from_source(self) -> None:
        source = """广西外国语学院期末成绩登记表
2025-2026学年第1学期
开课部门：E02软件工程系\t班级：软工2406班\t任课教师：张海林\t学分：3.0
课程名称：服务器配置与管理\t课程性质：专业教育\t考核方式：考试\t填表日期：2026-03-05
学号\t姓名\t平时\t期中\t实验/在线\t期末\t总评\t备注\t学号\t姓名\t平时\t期中\t实验/在线\t期末\t总评\t备注
24000000001\t左一\t90\t\t\t90\t90\t\t24000000003\t右一\t80\t\t\t80\t80\t
24000000002\t左二\t91\t\t\t91\t91
教师："""
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.doc"
            source_path.write_bytes(_gb18030_hex_rtf(source))
            conn = MagicMock()
            conn.execute.return_value.fetchone.return_value = {"file_hash": "abc"}
            payload = {
                "export_payload": {
                    "schema_version": "gxufl-academic-grade-register-v2",
                    "structured": {
                        "students": [
                            {"student_number": "24000000001", "student_name": "左一"},
                            {"student_number": "24000000003", "student_name": "右一"},
                            {"student_number": "24000000002", "student_name": "左二"},
                        ]
                    },
                }
            }
            with patch(
                "classroom_app.services.academic_final_material_service.resolve_global_file_path",
                return_value=source_path,
            ):
                repaired = repair_legacy_grade_register_roster_order(
                    conn,
                    {"document_type": ACADEMIC_GRADE_REGISTER_TYPE, "source_material_id": 99},
                    payload,
                )

        self.assertIn("FROM course_materials", conn.execute.call_args.args[0])
        export_payload = repaired["export_payload"]
        self.assertEqual("gxufl-academic-grade-register-v3", export_payload["schema_version"])
        self.assertEqual(
            ["24000000001", "24000000002", "24000000003"],
            [item["student_number"] for item in export_payload["structured"]["students"]],
        )

    def test_grade_register_rejects_rosters_larger_than_official_template_capacity(self) -> None:
        payload = build_grade_register_export_payload(self.grade, self.validation)
        payload["structured"]["students"] = [
            {
                "student_number": f"24000000{index + 1:03d}",
                "student_name": f"学生{index + 1:02d}",
                "final_score": 80,
            }
            for index in range(71)
        ]
        with self.assertRaisesRegex(ValueError, "最多容纳 70 名学生"):
            build_grade_register_docx(payload)

    def test_grade_register_signature_uses_the_official_floating_overlay(self) -> None:
        payload = build_grade_register_export_payload(self.grade, self.validation)
        with TemporaryDirectory() as temp_dir:
            signature_path = Path(temp_dir) / "signature.png"
            Image.new("RGBA", (160, 48), (0, 0, 0, 0)).save(signature_path)
            payload["fields"]["teacher_signature_image_path"] = str(signature_path)
            content = build_grade_register_docx(payload)

        with ZipFile(BytesIO(content)) as archive:
            self.assertIn("word/media/image1.png", archive.namelist())
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("<w:pict", document_xml)
        self.assertIn("margin-left:18.4pt", document_xml)
        self.assertIn("width:60.55pt", document_xml)
        self.assertIn("height:21.2pt", document_xml)

    def test_shared_export_contract_supports_both_types(self) -> None:
        grade_payload = build_grade_register_export_payload(self.grade, self.validation)
        analysis_payload = build_exam_analysis_export_payload(self.analysis, self.validation)
        for payload, expected_type in (
            (grade_payload, ACADEMIC_GRADE_REGISTER_TYPE),
            (analysis_payload, ACADEMIC_EXAM_ANALYSIS_TYPE),
        ):
            type_meta = resolve_material_ai_import_type("final_material", expected_type)
            self.assertEqual(expected_type, type_meta["key"])
            artifact = build_material_export_artifact(
                {"export_payload": payload},
                fallback_filename="academic-final",
                requested_format="docx",
            )
            self.assertTrue(artifact.content.startswith(b"PK"))
            self.assertTrue(artifact.filename.endswith(".docx"))

    def test_navigation_exposes_dedicated_final_material_group(self) -> None:
        nav = build_manage_nav(
            {"id": 1, "role": "teacher"},
            "academic_grade_registers",
            is_super_admin=False,
        )
        groups = {
            group["label"]: group
            for domain in nav["domains"]
            for group in domain["groups"]
        }
        self.assertIn("期末材料", groups)
        keys = {item["key"] for item in groups["期末材料"]["items"]}
        self.assertEqual({"academic_grade_registers", "academic_exam_analyses"}, keys)

    def test_schema_is_idempotent_and_enforces_one_batch_per_class(self) -> None:
        schema_academic_final_materials._SCHEMA_READY = False
        conn = sqlite3.connect(":memory:")
        try:
            schema_academic_final_materials.ensure_academic_final_material_schema(conn)
            schema_academic_final_materials.ensure_academic_final_material_schema(conn)
            columns = {
                str(row[1])
                for row in conn.execute("PRAGMA table_info(academic_final_material_batches)").fetchall()
            }
            self.assertIn("sync_options_json", columns)
            conn.execute(
                """
                INSERT INTO academic_final_material_batches
                (id, teacher_id, class_offering_id)
                VALUES ('one', 1, 8)
                """
            )
            with self.assertRaises(sqlite3.IntegrityError):
                conn.execute(
                    """
                    INSERT INTO academic_final_material_batches
                    (id, teacher_id, class_offering_id)
                    VALUES ('two', 1, 8)
                    """
                )
        finally:
            conn.close()
            schema_academic_final_materials._SCHEMA_READY = False

    def test_batch_upsert_is_idempotent_and_preserves_unspecified_state(self) -> None:
        schema_academic_final_materials._SCHEMA_READY = False
        conn = sqlite3.connect(":memory:")
        conn.row_factory = sqlite3.Row
        try:
            first = upsert_batch_state(
                conn,
                teacher_id=3,
                class_offering_id=9,
                values={
                    "sync_status": "running",
                    "course_name": "计算机网络",
                    "sync_options_json": json.dumps(
                        {"candidates": [{"exam_course_key": "course-1"}]},
                        ensure_ascii=False,
                    ),
                },
            )
            second = upsert_batch_state(
                conn,
                teacher_id=3,
                class_offering_id=9,
                values={"sync_status": "completed", "validation_status": "passed"},
            )
            self.assertEqual(first["id"], second["id"])
            self.assertEqual("计算机网络", second["course_name"])
            self.assertEqual("completed", second["sync_status"])
            self.assertEqual("course-1", second["sync_options"]["candidates"][0]["exam_course_key"])
            self.assertEqual(1, conn.execute("SELECT COUNT(*) FROM academic_final_material_batches").fetchone()[0])
        finally:
            conn.close()
            schema_academic_final_materials._SCHEMA_READY = False

    def test_stale_background_sync_becomes_retryable_failure(self) -> None:
        schema_academic_final_materials._SCHEMA_READY = False
        conn = sqlite3.connect(":memory:")
        conn.row_factory = sqlite3.Row
        try:
            schema_academic_final_materials.ensure_academic_final_material_schema(conn)
            old = (datetime.now() - timedelta(hours=1)).isoformat(timespec="seconds")
            conn.execute(
                """
                INSERT INTO academic_final_material_batches
                    (id, teacher_id, class_offering_id, sync_status, created_at, updated_at)
                VALUES ('stale', 7, 11, 'running', ?, ?)
                """,
                (old, old),
            )
            reclaimed = reclaim_stale_academic_final_material_batches(conn, 7)
            row = conn.execute(
                "SELECT sync_status, last_error FROM academic_final_material_batches WHERE id = 'stale'"
            ).fetchone()
            self.assertEqual(1, reclaimed)
            self.assertEqual("failed", row["sync_status"])
            self.assertIn("重新同步", row["last_error"])
        finally:
            conn.close()
            schema_academic_final_materials._SCHEMA_READY = False

    def test_candidate_query_does_not_mix_postgres_timestamp_and_text(self) -> None:
        class EmptyCursor:
            @staticmethod
            def fetchall():
                return []

        class RecordingConnection:
            def __init__(self):
                self.sql = ""

            def execute(self, sql, _params=()):
                self.sql = sql
                return EmptyCursor()

        conn = RecordingConnection()
        with patch(
            "classroom_app.services.academic_final_material_service.ensure_academic_final_material_schema"
        ):
            self.assertEqual([], list_teacher_final_material_candidates(conn, 7))

        self.assertNotIn("COALESCE(b.synced_at, o.created_at)", conn.sql)
        self.assertIn("(b.synced_at IS NULL) ASC", conn.sql)
        self.assertIn("b.synced_at DESC", conn.sql)
        self.assertIn("o.created_at DESC", conn.sql)

    def test_page_hides_state_containers_and_omits_redundant_notice(self) -> None:
        css = Path("static/css/academic_final_materials.css").read_text(encoding="utf-8")
        template = Path("templates/manage/academic_final_materials.html").read_text(encoding="utf-8")
        script = Path("static/js/academic_final_materials.js").read_text(encoding="utf-8")
        self.assertIn(".afm [hidden], .afm-dialog [hidden] { display: none !important; }", css)
        self.assertIn(".afm,\n.afm-dialog {", css)
        self.assertIn(".afm-dialog .afm-btn--primary:disabled", css)
        self.assertNotIn("一次同步，两张表共同更新", template)
        self.assertNotIn("正在读取已同步课程", template)
        self.assertIn("activeSyncStatuses", script)
        self.assertIn("schedulePolling", script)


class AcademicFinalMaterialBackgroundTaskTests(unittest.IsolatedAsyncioTestCase):
    async def test_ai_review_serializes_decimal_course_context(self) -> None:
        from classroom_app.routers.materials_parts import academic_final_materials as router_module

        ai_response = {
            "analysis_text": "本课程试题结构覆盖核心知识、综合应用与实践能力，难度梯度合理。成绩分布表明多数学生掌握扎实，少数学生在综合分析和知识迁移方面仍需加强。后续将增加分层案例、限时诊断和课堂复盘，并通过专项训练与过程反馈形成持续改进闭环。",
            "warnings": [],
        }
        with patch.object(router_module, "_call_ai_chat", return_value=ai_response) as ai_call:
            text, warnings, ai_used = await router_module._ai_review_and_analysis(
                {"fields": {"course_name": "动态Web程序设计"}},
                {"fields": {}, "structured": {}},
                {"credits": Decimal("2.0"), "assignments": []},
            )

        self.assertTrue(ai_used)
        self.assertFalse(warnings)
        self.assertGreaterEqual(len(text), 80)
        self.assertIn('"credits": 2.0', ai_call.call_args.args[1])

    async def test_ai_metadata_assistance_is_private_and_requires_deterministic_recheck(self) -> None:
        from classroom_app.routers.materials_parts import academic_final_materials as router_module

        grade = parse_grade_register_rtf(_rtf_bytes(GRADE_TEXT))
        analysis = parse_exam_analysis_rtf(_rtf_bytes(ANALYSIS_TEXT))
        context = {"course_name": "计算机网络实训", "teacher_name": "张老师"}
        validation = validate_paired_reports(grade, analysis, context=context, remote_student_count=2)
        result = {
            "grade": grade,
            "analysis": analysis,
            "context": context,
            "remote_students": [{}, {}],
            "validation": validation,
        }
        ai_response = {
            "equivalences": [
                {
                    "key": "context_course_name",
                    "equivalent": True,
                    "confidence": 0.99,
                    "reason": "本地课堂使用简称",
                }
            ]
        }

        with patch.object(router_module, "_call_ai_chat", return_value=ai_response) as ai_call:
            revised, warnings, ai_used = await router_module._ai_assist_metadata_validation(result)

        self.assertTrue(ai_used)
        self.assertTrue(revised["passed"], revised["errors"])
        self.assertTrue(revised["ai_assistance"]["deterministic_recheck_passed"])
        self.assertTrue(warnings)
        prompt = ai_call.call_args.args[1]
        self.assertNotIn("2300000001", prompt)
        self.assertNotIn("学生甲", prompt)

    async def test_ai_metadata_assistance_never_handles_score_failures(self) -> None:
        from classroom_app.routers.materials_parts import academic_final_materials as router_module

        grade = parse_grade_register_rtf(_rtf_bytes(GRADE_TEXT))
        analysis = parse_exam_analysis_rtf(_rtf_bytes(ANALYSIS_TEXT))
        analysis["statistics"]["average"] = 99.99
        validation = validate_paired_reports(grade, analysis, remote_student_count=2)
        result = {
            "grade": grade,
            "analysis": analysis,
            "context": {},
            "remote_students": [{}, {}],
            "validation": validation,
        }

        with patch.object(router_module, "_call_ai_chat") as ai_call:
            revised, warnings, ai_used = await router_module._ai_assist_metadata_validation(result)

        self.assertFalse(ai_used)
        self.assertFalse(revised["passed"])
        self.assertEqual([], warnings)
        ai_call.assert_not_called()

    async def test_report_download_sends_function_code_in_query_and_rejects_permission_page(self) -> None:
        class PermissionDeniedClient:
            base_url = "https://jwxt.gxufl.com"

            def __init__(self):
                self.params = {}
                self.get_called = False

            async def post(self, path, *, params, data, headers):
                self.params = dict(params)
                request = httpx.Request("POST", f"https://jwxt.gxufl.com{path}")
                return httpx.Response(200, request=request, text="<html>无功能权限</html>")

            async def get(self, *_args, **_kwargs):
                self.get_called = True
                raise AssertionError("权限页不得回退到公开空白报表")

        client = PermissionDeniedClient()
        with self.assertRaisesRegex(ValueError, "缺少报表权限"):
            await _download_fine_report_word(
                client,
                report_id="cjddy_bj.cpt",
                teaching_class_id="teaching-class-1",
                teacher_org_id="",
                source_summary=[],
            )
        self.assertEqual("N302505", client.params["gnmkdmKey"])
        self.assertFalse(client.get_called)

    async def test_report_download_posts_authorized_report_form_with_port(self) -> None:
        class AuthorizedReportClient:
            base_url = "https://jwxt.gxufl.com"

            def __init__(self):
                self.calls = []

            async def post(self, path, **kwargs):
                self.calls.append((str(path), kwargs))
                if str(path).startswith("/report/"):
                    request = httpx.Request("POST", f"https://jwxt.gxufl.com{path}")
                    return httpx.Response(
                        200,
                        request=request,
                        text="""
                            <form action="https://jwcjcx.gxufl.com:443/WebReport/ReportServer?reportlet=cjddy_bj.cpt"
                                  id="reportSearchForm" method="post">
                                <input type="hidden" name="jxb_id" value="teaching-class-1">
                            </form>
                        """,
                    )
                request = httpx.Request("POST", str(path))
                return httpx.Response(200, request=request, content=b"{\\rtf1 paired report}")

            async def get(self, *_args, **_kwargs):
                raise AssertionError("带参数的教务报表表单必须使用 POST")

        client = AuthorizedReportClient()
        content = await _download_fine_report_word(
            client,
            report_id="cjddy_bj.cpt",
            teaching_class_id="teaching-class-1",
            teacher_org_id="",
            source_summary=[],
        )
        self.assertEqual(b"{\\rtf1 paired report}", content)
        self.assertEqual(2, len(client.calls))
        self.assertEqual(
            "https://jwcjcx.gxufl.com:443/WebReport/ReportServer?reportlet=cjddy_bj.cpt",
            client.calls[1][0],
        )
        self.assertEqual({"jxb_id": "teaching-class-1"}, client.calls[1][1]["data"])
        self.assertTrue(client.calls[1][1]["follow_redirects"])

    async def test_sync_route_returns_queued_without_waiting_for_worker(self) -> None:
        from classroom_app.routers.materials_parts import academic_final_materials as router_module

        body = router_module.AcademicFinalMaterialSyncRequest(
            class_offering_id=81,
            force=True,
        )
        connection = MagicMock()
        connection_context = MagicMock()
        connection_context.__enter__.return_value = connection
        connection_context.__exit__.return_value = False
        queued_batch = {
            "id": "queued-batch",
            "teacher_id": 5,
            "class_offering_id": 81,
            "sync_status": "queued",
            "updated_at": datetime(2026, 8, 2, 4, 0, 0),
        }

        with (
            patch.object(router_module, "get_db_connection", return_value=connection_context),
            patch.object(router_module, "upsert_batch_state", return_value=queued_batch),
            patch.object(router_module, "_schedule_academic_final_material_sync", return_value=True),
        ):
            response = await router_module.api_sync_academic_final_materials(
                body,
                {"id": 5},
            )

        self.assertEqual(202, response.status_code)
        payload = json.loads(response.body)
        self.assertEqual("queued", payload["status"])
        self.assertEqual("queued", payload["batch"]["sync_status"])
        self.assertEqual("2026-08-02T04:00:00", payload["batch"]["updated_at"])
        connection.commit.assert_called_once_with()

    async def test_duplicate_background_schedules_are_coalesced(self) -> None:
        from classroom_app.routers.materials_parts import academic_final_materials as router_module

        started = asyncio.Event()
        release = asyncio.Event()

        async def fake_run(_body, _user):
            started.set()
            await release.wait()
            return {"status": "success"}

        body = router_module.AcademicFinalMaterialSyncRequest(class_offering_id=81)
        key = router_module._academic_final_material_task_key(5, 81)
        router_module._academic_final_material_tasks.pop(key, None)
        with patch.object(router_module, "_run_academic_final_material_sync", side_effect=fake_run):
            self.assertTrue(router_module._schedule_academic_final_material_sync(body, {"id": 5}))
            await asyncio.wait_for(started.wait(), timeout=1)
            self.assertFalse(router_module._schedule_academic_final_material_sync(body, {"id": 5}))
            release.set()
            task = router_module._academic_final_material_tasks[key]
            await asyncio.wait_for(asyncio.shield(task), timeout=1)
            await asyncio.sleep(0)
        self.assertNotIn(key, router_module._academic_final_material_tasks)


if __name__ == "__main__":
    unittest.main()
