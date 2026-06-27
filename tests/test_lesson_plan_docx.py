"""Tests for the GXUFL lesson-plan DOCX builder and Markdown parser."""

import unittest
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from classroom_app.services import lesson_plan_docx_service as docx_svc
from classroom_app.services import lesson_plan_markdown as md


_PROCESS_MD = """\
## 一、教学导入（8分钟）
回顾上节课内容。**高阶提问**：为什么需要规范构建环境？

## 二、讲授新课
| 教学环节 | 教学活动（教师引导） | 学生活动（主体） | 设计意图（OBE & 两性一度） |
| :--- | :--- | :--- | :--- |
| 任务1 | 手把手演示 Maven 环境配置 | 动手实践 | 高阶性：建立工程化思维 |
| 任务2 | 设置端口冲突排错情境 | 自主排错 | 挑战度：训练日志定位 |

## 三、教学小结
- 梳理知识
- 布置作业
"""

_PLAN = {
    "cover": {
        "course_name": "动态web程序设计",
        "course_category": "专业任选课程",
        "credits": "2.0",
        "total_hours": "32",
        "teacher_name": "张海林",
        "teaching_unit": "数字科技学院",
        "class_name": "软工2401班",
        "textbook": "《Spring+Spring MVC+MyBatis+Spring Boot框架整合开发》",
        "publisher": "人民邮电出版社",
        "semester_label": "2025—2026学年第二学期",
        "school_name": "广西外国语学院",
    },
    "sessions": [
        {
            "index": 1,
            "schedule": {"text": "2026年 03月 09 日 第 一 周 星期 一  第 10-11 节"},
            "chapter": "第一讲：开宗明义——开发环境构建与Spring初探",
            "objectives": "知识目标：了解 Spring 生态。\n能力目标：能够完成 Maven 配置。",
            "key_points": "Maven 核心配置；Spring Boot 项目创建流程。",
            "difficulties": "环境变量与路径排错。",
            "methods": "讲授法，演示法，案例法",
            "means": "课件，上机实操",
            "process": _PROCESS_MD,
            "side_notes": "提醒重开终端以生效变量",
            "post_notes": "类比教学效果较好。",
        },
        {
            "index": 2,
            "schedule": {"text": "2026年 03月 16 日 第 二 周 星期 一  第 10-11 节"},
            "chapter": "第二讲：Spring Boot 控制器",
            "objectives": "掌握 Controller 的基本写法。",
            "key_points": "注解与请求映射。",
            "difficulties": "参数绑定。",
            "methods": "任务驱动",
            "means": "PPT，上机实操",
            "process": "## 一、教学导入\n情境创设。",
            "side_notes": "",
            "post_notes": "",
        },
    ],
}


def _tbl_grid_widths(table) -> list[int]:
    grid = table._tbl.find(qn("w:tblGrid"))
    return [int(col.get(qn("w:w"))) for col in grid.findall(qn("w:gridCol"))]


class MarkdownParserTests(unittest.TestCase):
    def test_parses_heading_table_and_list(self):
        blocks = md.parse_blocks(_PROCESS_MD)
        types = [b["type"] for b in blocks]
        self.assertIn("heading", types)
        self.assertIn("table", types)
        self.assertIn("ul", types)
        table = next(b for b in blocks if b["type"] == "table")
        self.assertEqual(len(table["header"]), 4)
        self.assertEqual(len(table["rows"]), 2)

    def test_inline_bold(self):
        runs = md.inline_runs("普通**加粗**普通")
        self.assertTrue(any(r["bold"] and r["text"] == "加粗" for r in runs))

    def test_html_escapes(self):
        html_out = md.markdown_to_html("a < b & c")
        self.assertIn("&lt;", html_out)
        self.assertIn("&amp;", html_out)


class DocxBuilderTests(unittest.TestCase):
    def setUp(self):
        self.doc_bytes = docx_svc.build_lesson_plan_docx(_PLAN)
        self.document = Document(BytesIO(self.doc_bytes))

    def test_returns_bytes(self):
        self.assertIsInstance(self.doc_bytes, bytes)
        self.assertGreater(len(self.doc_bytes), 1000)

    def test_uses_reference_page_setup(self):
        section = self.document.sections[0]
        self.assertEqual(section.page_width, 7_560_310)
        self.assertEqual(section.page_height, 10_692_130)
        self.assertEqual(section.top_margin, 914_400)
        self.assertEqual(section.left_margin, 810_260)

    def test_cover_is_not_a_grid_table(self):
        cover_text = "\n".join(p.text for p in self.document.paragraphs)
        self.assertIn("教  案", cover_text)
        self.assertIn("动态web程序设计", cover_text)
        self.assertIn("2025—2026学年第 二 学期", cover_text)
        # Top-level tables are only session tables; the school cover uses
        # underlined paragraphs, not the old gray cover table.
        self.assertEqual(len(self.document.tables), len(_PLAN["sessions"]))

    def test_cover_uses_first_page_footer_for_imprint(self):
        footer_text = "\n".join(p.text for p in self.document.sections[0].first_page_footer.paragraphs)
        body_text = "\n".join(p.text for p in self.document.paragraphs)
        self.assertIn("广西外国语学院教务处 印制", footer_text)
        self.assertNotIn("广西外国语学院教务处 印制", body_text)

    def test_long_textbook_continuation_starts_at_field_left(self):
        plan = deepcopy(_PLAN)
        plan["cover"]["textbook"] = "《Spring+Spring MVC+MyBatis+Spring Boot框架整合开发（IntelliJ IDEA版·微课视频版）》"
        plan["cover"]["publisher"] = "人民邮电出版社"
        document = Document(BytesIO(docx_svc.build_lesson_plan_docx(plan)))
        publisher_para = next(p for p in document.paragraphs if "人民邮电出版社" in p.text)
        publisher_run = next(r.text for r in publisher_para.runs if "人民邮电出版社" in r.text)
        self.assertTrue(publisher_run.startswith("人民邮电出版社"))

    def test_semester_year_and_term_are_underlined(self):
        semester_para = next(p for p in self.document.paragraphs if "学年第" in p.text)
        self.assertTrue(any(r.text == "2025—2026" and r.underline for r in semester_para.runs))
        self.assertTrue(any(r.text == "二" and r.underline for r in semester_para.runs))

    def test_sessions_are_page_separated_without_extra_caption(self):
        breaks = self.document.element.body.findall(".//" + qn("w:br"))
        page_breaks = [br for br in breaks if br.get(qn("w:type")) == "page"]
        self.assertEqual(page_breaks, [])
        paragraph_breaks = self.document.element.body.findall(".//" + qn("w:pageBreakBefore"))
        self.assertGreaterEqual(len(paragraph_breaks), len(_PLAN["sessions"]))
        all_paragraphs = "\n".join(p.text for p in self.document.paragraphs)
        self.assertNotIn("第 1 次课", all_paragraphs)

    def test_session_table_uses_reference_grid(self):
        self.assertEqual(_tbl_grid_widths(self.document.tables[0]), [1911, 324, 6095, 1678])
        first_row = self.document.tables[0].rows[0]
        self.assertIn("授课时间", first_row.cells[0].text)
        self.assertIn("2026 年 03 月 09 日", first_row.cells[2].text)
        runs = first_row.cells[2].paragraphs[0].runs
        self.assertTrue(any(r.text == "03" and r.underline for r in runs))
        self.assertTrue(any(r.text == "10-11" and r.underline for r in runs))

    def test_objective_and_key_labels_are_bold(self):
        table = self.document.tables[0]
        objective_runs = table.rows[2].cells[2].paragraphs[0].runs
        key_runs = table.rows[3].cells[2].paragraphs[0].runs
        difficult_runs = table.rows[3].cells[2].paragraphs[1].runs
        self.assertTrue(any(r.text == "知识目标：" and r.bold for r in objective_runs))
        self.assertTrue(any(r.text == "重点：" and r.bold for r in key_runs))
        self.assertTrue(any(r.text == "难点：" and r.bold for r in difficult_runs))

    def test_markdown_table_rendered_as_reference_nested_table(self):
        nested = []
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    nested.extend(cell.tables)
        self.assertTrue(nested, "PBL markdown table not rendered as nested docx table")
        activity = nested[0]
        self.assertEqual(_tbl_grid_widths(activity), [1141, 2684, 2439, 1844])
        header = " ".join(c.text for c in activity.rows[0].cells)
        self.assertIn("教学环节", header)
        self.assertIn("设计意图", header)

    def test_session_table_has_required_labels(self):
        joined = "\n".join(c.text for t in self.document.tables for r in t.rows for c in r.cells)
        for label in (
            "授课时间",
            "授课章节",
            "教学目的和要求",
            "教学重点和难点",
            "教学方法和手段",
            "教学内容及过程",
            "旁批",
            "教学后记",
        ):
            self.assertIn(label, joined, f"missing label {label}")

    def test_mermaid_code_block_is_inserted_as_image(self):
        plan = deepcopy(_PLAN)
        plan["sessions"] = [
            {
                **plan["sessions"][0],
                "process": "```mermaid\nflowchart TD\nA[开始] --> B[处理]\nB --> C[结束]\n```",
            }
        ]
        document = Document(BytesIO(docx_svc.build_lesson_plan_docx(plan)))
        self.assertGreaterEqual(len(document.inline_shapes), 2)
        joined = "\n".join(c.text for t in document.tables for r in t.rows for c in r.cells)
        self.assertNotIn("flowchart TD", joined)


if __name__ == "__main__":
    unittest.main()
