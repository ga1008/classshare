"""Official DOCX renderers for JWXT-downloaded final materials."""

from __future__ import annotations

import hashlib
import io
import math
import re
from pathlib import Path
from typing import Any

from .academic_final_material_service import (
    ACADEMIC_EXAM_ANALYSIS_LABEL,
    ACADEMIC_EXAM_ANALYSIS_TYPE,
    ACADEMIC_GRADE_REGISTER_LABEL,
    ACADEMIC_GRADE_REGISTER_TYPE,
)


_GRADE_REGISTER_TEMPLATE_PATH = Path(__file__).with_name("assets") / "gxufl_academic_grade_register_template.docx"
_GRADE_REGISTER_TEMPLATE_SHA256 = "8351a5027010c6f38e3edaeb6b49252d691e1d0a22c0f301a760abeb742d2b9e"
_GRADE_REGISTER_LEFT_CAPACITY = 40
_GRADE_REGISTER_RIGHT_CAPACITY = 30
_GRADE_REGISTER_MAX_STUDENTS = _GRADE_REGISTER_LEFT_CAPACITY + _GRADE_REGISTER_RIGHT_CAPACITY
_GRADE_REGISTER_GRID_WIDTHS = (
    1473, 850, 566, 510, 566, 567, 567, 511, 396,
    1078, 851, 567, 511, 567, 567, 567, 511,
)
_GRADE_REGISTER_ROW_HEIGHTS = (
    480,
    280,
    340,
    340,
    500,
    *([340] * 32),
    500,
    *([340] * 7),
    180,
    280,
)
_GRADE_REGISTER_STUDENT_KEYS = (
    "student_number",
    "student_name",
    "ordinary_score",
    "midterm_score",
    "experiment_online_score",
    "final_exam_score",
    "final_score",
    "remark",
)
_GRADE_REGISTER_LEFT_COLUMNS = (0, 1, 2, 3, 4, 5, 6, 7)
_GRADE_REGISTER_RIGHT_COLUMNS = (8, 10, 11, 12, 13, 14, 15, 16)


def _payload(parse_payload: dict[str, Any]) -> dict[str, Any]:
    export_payload = parse_payload.get("export_payload")
    return export_payload if isinstance(export_payload, dict) else parse_payload


def _safe_filename(value: Any) -> str:
    return re.sub(r'[\\/:*?"<>|\r\n]+', "_", str(value or "")).strip(" ._")[:100] or "期末材料"


def build_academic_final_material_filename(parse_payload: dict[str, Any], template_key: str) -> str:
    payload = _payload(parse_payload)
    fields = payload.get("fields") if isinstance(payload.get("fields"), dict) else {}
    label = ACADEMIC_GRADE_REGISTER_LABEL if template_key == ACADEMIC_GRADE_REGISTER_TYPE else ACADEMIC_EXAM_ANALYSIS_LABEL
    course = _safe_filename(fields.get("course_name") or "课程")
    class_name = _safe_filename(fields.get("class_name") or "班级")
    return f"{label}-{course}-{class_name}.docx"


def _set_run_font(run: Any, size: float, *, bold: bool = False, name: str = "宋体") -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _set_cell_text(
    cell: Any,
    value: Any,
    *,
    size: float = 8.5,
    bold: bool = False,
    align: int = 1,
    name: str = "宋体",
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("" if value is None else str(value))
    _set_run_font(run, size, bold=bold, name=name)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_width(cell: Any, width_cm: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_borders(table: Any, *, size: int = 6) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)

def _add_title(document: Any, title: str, fields: dict[str, Any], *, size: float = 16) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = Pt(2)
    _set_run_font(paragraph.add_run(title), size, bold=True)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_before = 0
    period.paragraph_format.space_after = Pt(3)
    academic_year = str(fields.get("academic_year") or "")
    semester = str(fields.get("semester") or "")
    _set_run_font(period.add_run(f"{academic_year}学年{semester}".strip()), 10.5, bold=True)


def _add_image_to_cell(
    cell: Any,
    path_value: Any,
    *,
    width_cm: float,
    height_cm: float | None = None,
    prefix: str = "",
) -> bool:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    path = Path(str(path_value or ""))
    if not path.is_file():
        return False
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    if prefix:
        run = paragraph.add_run(prefix)
        _set_run_font(run, 8)
    kwargs = {"width": Cm(width_cm)}
    if height_cm:
        kwargs["height"] = Cm(height_cm)
    paragraph.add_run().add_picture(str(path), **kwargs)
    return True


def _score_text(value: Any) -> str:
    if value in (None, ""):
        return ""
    parsed = _finite_float(value)
    if parsed is None:
        return str(value).strip()
    if abs(parsed - round(parsed)) < 0.001:
        return str(int(round(parsed)))
    return f"{parsed:.2f}"


def _finite_float(value: Any) -> float | None:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return None
    return parsed if math.isfinite(parsed) else None


def _set_template_cell_text(cell: Any, value: Any) -> None:
    from docx.oxml.ns import qn

    text_nodes = list(cell._tc.iter(qn("w:t")))
    text = "" if value is None else str(value)
    if not text_nodes:
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, 9, name="微软雅黑")
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


def _academic_period_text(fields: dict[str, Any]) -> str:
    academic_year = re.sub(r"\s+", "", str(fields.get("academic_year") or ""))
    semester = re.sub(r"\s+", "", str(fields.get("semester") or ""))
    if academic_year and not academic_year.endswith("学年"):
        academic_year = f"{academic_year}学年"
    normalized_semester = semester.replace("学期", "").removeprefix("第")
    semester_number = {
        "1": "1", "一": "1", "上": "1", "秋季": "1",
        "2": "2", "二": "2", "下": "2", "春季": "2",
    }.get(normalized_semester)
    if semester_number:
        semester = f"第{semester_number}学期"
    elif semester and not semester.endswith("学期"):
        semester = f"{semester}学期"
    return f"{academic_year}{semester}"


def _grade_student_value(student: dict[str, Any] | None, key: str) -> str:
    if not student:
        return ""
    value = student.get(key)
    if key == "final_score" and value not in (None, ""):
        parsed = _finite_float(value)
        if parsed is None:
            return str(value).strip()
        return "100" if abs(parsed - 100) < 0.001 else f"{parsed:.2f}"
    if key in {
        "ordinary_score",
        "midterm_score",
        "experiment_online_score",
        "final_exam_score",
    }:
        return _score_text(value)
    return "" if value is None else str(value)


def _score_band_counts(students: list[dict[str, Any]]) -> list[int]:
    counts = [0] * 8
    for student in students:
        score = _finite_float(student.get("final_exam_score"))
        if score is None:
            continue
        if score >= 90:
            index = 0
        elif score >= 80:
            index = 1
        elif score >= 70:
            index = 2
        elif score >= 60:
            index = 3
        elif score >= 50:
            index = 4
        elif score >= 40:
            index = 5
        elif score >= 30:
            index = 6
        else:
            index = 7
        counts[index] += 1
    return counts


def _load_verified_grade_register_template() -> Any:
    """Load only the audited, privacy-scrubbed official template asset."""
    from docx import Document
    from docx.oxml.ns import qn

    if not _GRADE_REGISTER_TEMPLATE_PATH.is_file():
        raise RuntimeError("期末成绩登记表官方版式模板缺失，无法导出。")
    content = _GRADE_REGISTER_TEMPLATE_PATH.read_bytes()
    actual_hash = hashlib.sha256(content).hexdigest()
    if actual_hash != _GRADE_REGISTER_TEMPLATE_SHA256:
        raise RuntimeError("期末成绩登记表官方版式模板校验失败，请恢复经过像素验收的模板资产。")
    document = Document(io.BytesIO(content))
    if len(document.sections) != 1 or len(document.tables) != 1:
        raise RuntimeError("期末成绩登记表官方版式模板结构已损坏。")
    section = document.sections[0]
    section_geometry = (
        section.page_width.twips,
        section.page_height.twips,
        section.top_margin.twips,
        section.right_margin.twips,
        section.bottom_margin.twips,
        section.left_margin.twips,
    )
    if section_geometry != (11905, 16837, 226, 283, 5, 283):
        raise RuntimeError("期末成绩登记表官方版式模板页面参数已漂移。")
    table = document.tables[0]
    if len(table.rows) != 47 or len(table.columns) != 17:
        raise RuntimeError("期末成绩登记表官方版式模板表格结构已损坏。")
    grid_widths = tuple(int(node.get(qn("w:w"))) for node in table._tbl.tblGrid)
    row_heights = tuple(
        int(row._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")))
        for row in table.rows
    )
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.find(qn("w:tblW"))
    table_indent = tbl_pr.find(qn("w:tblInd"))
    table_layout = tbl_pr.find(qn("w:tblLayout"))
    if (
        grid_widths != _GRADE_REGISTER_GRID_WIDTHS
        or row_heights != _GRADE_REGISTER_ROW_HEIGHTS
        or table_width is None
        or table_width.get(qn("w:w")) != "11225"
        or table_indent is None
        or table_indent.get(qn("w:w")) != "5"
        or table_layout is None
        or table_layout.get(qn("w:type")) != "fixed"
    ):
        raise RuntimeError("期末成绩登记表官方版式模板表格参数已漂移。")
    return document


def _grade_status_counts(students: list[dict[str, Any]]) -> dict[str, int]:
    aliases = {
        "免考(修)": ("免考", "免修"),
        "缓考": ("缓考",),
        "作弊": ("作弊",),
        "旷考": ("旷考",),
        "交流生": ("交流生",),
        "借读生": ("借读生",),
    }
    counts = {key: 0 for key in aliases}
    for student in students:
        remark = str(student.get("remark") or "").strip()
        for key, values in aliases.items():
            if any(value in remark for value in values):
                counts[key] += 1
                break
    return counts


def _add_grade_signature_overlay(document: Any, table: Any, signature_path: Any, signature_count: int = 1) -> None:
    path = Path(str(signature_path or ""))
    if not path.is_file():
        return

    from docx.oxml import parse_xml

    relationship_id, _image = document.part.get_or_add_image(str(path))
    width_pt = 60.55
    height_pt = 21.2
    if int(signature_count or 1) > 1:
        try:
            from PIL import Image

            with Image.open(path) as source:
                aspect_ratio = source.width / max(1, source.height)
            width_pt = min(130.0, max(60.55, height_pt * aspect_ratio))
        except Exception:
            width_pt = min(130.0, 60.55 * int(signature_count or 1))
    pict = parse_xml(
        f"""
        <w:pict
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            w14:anchorId="4E0449E0">
          <v:shapetype id="_x0000_t75" coordsize="21600,21600" o:spt="75"
              o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe" filled="f" stroked="f">
            <v:stroke joinstyle="miter"/>
            <v:formulas>
              <v:f eqn="if lineDrawn pixelLineWidth 0"/><v:f eqn="sum @0 1 0"/>
              <v:f eqn="sum 0 0 @1"/><v:f eqn="prod @2 1 2"/>
              <v:f eqn="prod @3 21600 pixelWidth"/><v:f eqn="prod @3 21600 pixelHeight"/>
              <v:f eqn="sum @0 0 1"/><v:f eqn="prod @6 1 2"/>
              <v:f eqn="prod @7 21600 pixelWidth"/><v:f eqn="sum @8 21600 0"/>
              <v:f eqn="prod @7 21600 pixelHeight"/><v:f eqn="sum @10 21600 0"/>
            </v:formulas>
            <v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>
            <o:lock v:ext="edit" aspectratio="t"/>
          </v:shapetype>
          <v:shape id="_x0000_s2050" type="#_x0000_t75"
              style="position:absolute;margin-left:18.4pt;margin-top:3.15pt;width:{width_pt:g}pt;height:{height_pt:g}pt;z-index:251659264;mso-position-horizontal-relative:text;mso-position-vertical-relative:text">
            <v:imagedata r:id="{relationship_id}" o:title=""/>
          </v:shape>
        </w:pict>
        """
    )
    anchor_paragraph = table.rows[45].cells[6].paragraphs[0]
    anchor_run = anchor_paragraph.runs[0] if anchor_paragraph.runs else anchor_paragraph.add_run()
    anchor_run._r.append(pict)


def build_grade_register_docx(parse_payload: dict[str, Any]) -> bytes:
    payload = _payload(parse_payload)
    fields = payload.get("fields") if isinstance(payload.get("fields"), dict) else {}
    structured = payload.get("structured") if isinstance(payload.get("structured"), dict) else {}
    students = structured.get("students") if isinstance(structured.get("students"), list) else []
    students = [student for student in students if isinstance(student, dict)]
    if len(students) > _GRADE_REGISTER_MAX_STUDENTS:
        raise ValueError(f"期末成绩登记表模板最多容纳 {_GRADE_REGISTER_MAX_STUDENTS} 名学生，当前为 {len(students)} 名。")
    document = _load_verified_grade_register_template()
    table = document.tables[0]

    metadata_cells = {
        (0, 0): "广西外国语学院期末成绩登记表",
        (1, 0): _academic_period_text(fields),
        (2, 0): f"开课部门：{fields.get('department') or ''}",
        (2, 4): f"班级：{fields.get('class_name') or ''}",
        (2, 10): f"任课教师：{fields.get('teacher_name') or ''}",
        (2, 14): f"学分：{fields.get('credits') or ''}",
        (3, 0): f"课程名称：{fields.get('course_name') or ''}",
        (3, 4): f"课程性质：{fields.get('course_nature') or ''}",
        (3, 8): f"考核方式：{fields.get('assessment_method') or ''}",
        (3, 11): f"填表日期：{fields.get('date') or ''}",
    }
    for (row_index, column_index), value in metadata_cells.items():
        _set_template_cell_text(table.rows[row_index].cells[column_index], value)

    padded_students: list[dict[str, Any] | None] = [*students]
    padded_students.extend([None] * (_GRADE_REGISTER_MAX_STUDENTS - len(padded_students)))
    for slot in range(_GRADE_REGISTER_LEFT_CAPACITY):
        row = table.rows[5 + slot]
        student = padded_students[slot]
        for key, column_index in zip(_GRADE_REGISTER_STUDENT_KEYS, _GRADE_REGISTER_LEFT_COLUMNS):
            _set_template_cell_text(row.cells[column_index], _grade_student_value(student, key))
    for slot in range(_GRADE_REGISTER_RIGHT_CAPACITY):
        row = table.rows[5 + slot]
        student = padded_students[_GRADE_REGISTER_LEFT_CAPACITY + slot]
        for key, column_index in zip(_GRADE_REGISTER_STUDENT_KEYS, _GRADE_REGISTER_RIGHT_COLUMNS):
            _set_template_cell_text(row.cells[column_index], _grade_student_value(student, key))

    formula = str(structured.get("formula") or "平时*40% + 期末*60%").strip()
    _set_template_cell_text(table.rows[35].cells[9], f"总评成绩 = {formula}")

    band_counts = _score_band_counts(students)
    total = len(students)
    for offset, count in enumerate(band_counts):
        row = table.rows[37 + offset]
        _set_template_cell_text(row.cells[10], count)
        ratio = count * 100 / total if total else 0
        _set_template_cell_text(row.cells[12], "0%" if count == 0 else f"{ratio:.2f}%")

    status_counts = _grade_status_counts(students)
    statistics = structured.get("statistics") if isinstance(structured.get("statistics"), dict) else {}
    numeric_scores = [
        score
        for student in students
        if (score := _finite_float(student.get("final_exam_score"))) is not None
    ]
    examined = len(numeric_scores)
    average = _finite_float(statistics.get("average"))
    if average is None:
        average = sum(numeric_scores) / len(numeric_scores) if numeric_scores else 0.0
    summary_values = [
        f"{total}/{examined}",
        status_counts["免考(修)"],
        status_counts["缓考"],
        status_counts["作弊"],
        status_counts["旷考"],
        status_counts["交流生"],
        status_counts["借读生"],
        f"{average:.2f}",
    ]
    for offset, value in enumerate(summary_values):
        _set_template_cell_text(table.rows[37 + offset].cells[16], value)

    _set_template_cell_text(table.rows[46].cells[0], "教师：______________________________签字")
    teacher_signature_ids = fields.get("teacher_signature_ids") if isinstance(fields.get("teacher_signature_ids"), list) else []
    _add_grade_signature_overlay(
        document,
        table,
        fields.get("teacher_signature_image_path"),
        len(teacher_signature_ids) or (1 if fields.get("teacher_signature_image_path") else 0),
    )
    if "{{" in document._element.xml:
        raise RuntimeError("期末成绩登记表仍包含未替换的模板占位符。")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _choice(options: list[str], selected: Any) -> str:
    selected_text = str(selected or "").strip()
    return "   ".join(f"{'√' if option == selected_text else '□'} {option}" for option in options)


def _chart_image(distribution: list[dict[str, Any]]) -> io.BytesIO:
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1250, 390
    image = Image.new("RGB", (width, height), "#D0D0D0")
    draw = ImageDraw.Draw(image)
    font_paths = (
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    )
    font_path = next((path for path in font_paths if path.is_file()), None)
    label_font = ImageFont.truetype(str(font_path), 24) if font_path else ImageFont.load_default()
    value_font = ImageFont.truetype(str(font_path), 25) if font_path else ImageFont.load_default()
    axis_font = ImageFont.truetype(str(font_path), 21) if font_path else ImageFont.load_default()
    left, top, right, bottom = 88, 35, width - 28, height - 58
    draw.line((left, top, left, bottom), fill="#B2B2B2", width=2)
    draw.line((left, bottom, right, bottom), fill="#B2B2B2", width=2)
    counts = [int(item.get("count") or 0) for item in distribution[:5]]
    maximum = max(counts or [1]) or 1
    axis_max = max(5, ((maximum + 4) // 5) * 5)
    labels = ["<60", "60-69", "70-79", "80-89", "90-100"]
    colors = ["#A7A7A7", "#70AD47", "#ED7D31", "#FF3B3B", "#20A9C2"]
    for tick in range(0, axis_max + 1, 5):
        y = bottom - (bottom - top) * tick / axis_max
        if tick:
            draw.line((left, y, right, y), fill="#C6C6C6", width=1)
        label = str(tick)
        box = draw.textbbox((0, 0), label, font=axis_font)
        draw.text((left - 12 - (box[2] - box[0]), y - (box[3] - box[1]) / 2), label, fill="#4B4B4B", font=axis_font)
    slot = (right - left) / 5
    for index, count in enumerate(counts + [0] * (5 - len(counts))):
        bar_width = slot * 0.34
        x0 = left + index * slot + slot * 0.33
        x1 = x0 + bar_width
        bar_height = (bottom - top - 12) * count / axis_max
        y0 = bottom - bar_height
        draw.rectangle((x0, y0, x1, bottom), fill=colors[index])
        count_box = draw.textbbox((0, 0), str(count), font=value_font)
        draw.text(((x0 + x1 - (count_box[2] - count_box[0])) / 2, max(top, y0 - 30)), str(count), fill="#222222", font=value_font)
        label_box = draw.textbbox((0, 0), labels[index], font=label_font)
        draw.text(((x0 + x1 - (label_box[2] - label_box[0])) / 2, bottom + 12), labels[index], fill="#333333", font=label_font)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    buffer.seek(0)
    return buffer


def _merge_row(table: Any, row_index: int, start: int, end: int) -> Any:
    cell = table.rows[row_index].cells[start]
    if end > start:
        cell = cell.merge(table.rows[row_index].cells[end])
    return cell


_ANALYSIS_GRID_POINTS = (21.7, 46.5, 31.5, 54.0, 28.5, 43.5, 71.25, 45.75, 54.0, 42.0, 43.5)


def _set_cell_width_points(cell: Any, width_points: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.width = Pt(width_points)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_points * 20)))
    tc_w.set(qn("w:type"), "dxa")


def _set_row_height(row: Any, points: float | None, *, exact: bool = True) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if points is not None:
        row.height = Pt(points)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "1")
    tr_pr.append(cant_split)


def _set_cell_direction_vertical(cell: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    direction = OxmlElement("w:textDirection")
    direction.set(qn("w:val"), "tbRl")
    tc_pr.append(direction)


def _remove_cell_borders(cell: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _configure_analysis_table(table: Any) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, size=4)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.find(qn("w:tblW"))
    if width is not None:
        width.set(qn("w:w"), str(round(482.2 * 20)))
        width.set(qn("w:type"), "dxa")
    margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)
    grid_columns = list(table._tbl.tblGrid)
    for grid_column, points in zip(grid_columns, _ANALYSIS_GRID_POINTS):
        grid_column.set(qn("w:w"), str(round(points * 20)))
    for row in table.rows:
        for column, cell in enumerate(row.cells):
            _set_cell_width_points(cell, _ANALYSIS_GRID_POINTS[column])


def _analysis_cell(table: Any, row: int, start: int, end: int, text: Any = "", *, size: float = 9, bold: bool = False, align: int = 1) -> Any:
    cell = _merge_row(table, row, start, end)
    _set_cell_text(cell, text, size=size, bold=bold, align=align)
    return cell


def _check(selected: Any, option: str) -> str:
    return "√" if str(selected or "").strip() == option else ""


def _set_analysis_body(cell: Any, value: Any) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    lines = [line.strip() for line in str(value or "").replace("\r", "").split("\n") if line.strip()]
    if not lines:
        lines = [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.first_line_indent = Pt(24) if not re.match(r"^[一二三四五六七八九十]+[、.]|^\d+[.、]", line) else Pt(0)
        _set_run_font(paragraph.add_run(line), 12)


def _set_review_signature(cell: Any, path_value: Any, signature_count: int = 1) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    _set_run_font(paragraph.add_run("签字："), 9)
    path = Path(str(path_value or ""))
    if path.is_file():
        if int(signature_count or 1) <= 1:
            paragraph.add_run().add_picture(str(path), width=Cm(3.1), height=Cm(1.1))
        else:
            paragraph.add_run().add_picture(str(path), width=Cm(4.35))


def _build_analysis_table(document: Any, fields: dict[str, Any], structured: dict[str, Any]) -> Any:
    from docx.shared import Cm

    table = document.add_table(rows=22, cols=11)
    _configure_analysis_table(table)
    heights = (36, 17, 11, 14, 14, 14, 15, 14, 14, 15, 15, 15, 15, 14, 14, 150, 14, 248, 14, None, 77.1, 14)
    for row, height in zip(table.rows, heights):
        _set_row_height(row, height, exact=height is not None)

    title = _analysis_cell(table, 0, 0, 10, "广西外国语学院课程试卷分析表", size=16)
    period = _analysis_cell(table, 1, 0, 10, _academic_period_text(fields), size=12)
    spacer = _analysis_cell(table, 2, 0, 10, "", size=9)
    for cell in (title, period, spacer):
        _remove_cell_borders(cell)

    _analysis_cell(table, 3, 0, 1, "课程名称")
    _analysis_cell(table, 3, 2, 4, fields.get("course_name") or "", align=0)
    _analysis_cell(table, 3, 5, 5, "学时数")
    _analysis_cell(table, 3, 6, 6, fields.get("course_hours") or "")
    _analysis_cell(table, 3, 7, 7, "开课单位")
    _analysis_cell(table, 3, 8, 10, fields.get("department") or "")

    _analysis_cell(table, 4, 0, 1, "教师姓名")
    _analysis_cell(table, 4, 2, 4, fields.get("teacher_name") or "")
    _analysis_cell(table, 4, 5, 6, "课程性质")
    _analysis_cell(table, 4, 7, 7, "选修")
    _analysis_cell(table, 4, 8, 8, _check(fields.get("course_nature"), "选修"))
    _analysis_cell(table, 4, 9, 9, "必修")
    _analysis_cell(table, 4, 10, 10, _check(fields.get("course_nature"), "必修"))

    _analysis_cell(table, 5, 0, 1, "命题形式(打√)")
    proposition_cells = ((2, 2, "试题库"), (5, 5, "试卷库"), (7, 8, "教师组题"))
    check_cells = ((3, 4, "试题库"), (6, 6, "试卷库"), (9, 10, "教师组题"))
    for start, end, label in proposition_cells:
        _analysis_cell(table, 5, start, end, label)
    for start, end, label in check_cells:
        _analysis_cell(table, 5, start, end, _check(fields.get("proposition_form"), label))

    _analysis_cell(table, 6, 0, 1, "考试形式(打√)")
    for start, value in ((2, "开卷"), (4, "闭卷")):
        _analysis_cell(table, 6, start, start, value)
        _analysis_cell(table, 6, start + 1, start + 1, _check(fields.get("exam_form"), value))
    _analysis_cell(table, 6, 6, 6, "教考分离(打√)")
    _analysis_cell(table, 6, 7, 7, "是")
    _analysis_cell(table, 6, 8, 8, _check(fields.get("separate_teaching_exam"), "是"))
    _analysis_cell(table, 6, 9, 9, "否")
    _analysis_cell(table, 6, 10, 10, _check(fields.get("separate_teaching_exam"), "否"))

    _analysis_cell(table, 7, 0, 1, "学生班级")
    _analysis_cell(table, 7, 2, 10, fields.get("class_name") or "", align=0)

    distribution = structured.get("score_distribution") if isinstance(structured.get("score_distribution"), list) else []
    distribution = [item for item in distribution if isinstance(item, dict)]
    while len(distribution) < 5:
        distribution.append({"count": 0, "ratio": 0})
    vertical = table.cell(8, 0).merge(table.cell(12, 0))
    _set_cell_direction_vertical(vertical)
    _set_cell_text(vertical, "分数分布", size=9, bold=True)
    row_specs = ((2, 3), (4, 5), (6, 6), (7, 8), (9, 10))
    for row_index, label, values in (
        (8, "分数段", ["<60", "60-69", "70-79", "80-89", "90-100"]),
        (9, "人数", [int(item.get("count") or 0) for item in distribution[:5]]),
        (10, "比例", [f"{float(item.get('ratio') or 0):.2f}%" for item in distribution[:5]]),
    ):
        _analysis_cell(table, row_index, 1, 1, label)
        for (start, end), value in zip(row_specs, values):
            _analysis_cell(table, row_index, start, end, value)
    stats = structured.get("statistics") if isinstance(structured.get("statistics"), dict) else {}
    _analysis_cell(table, 11, 1, 1, "平均分")
    _analysis_cell(table, 11, 2, 5, f"{float(stats.get('average') or 0):.2f}")
    _analysis_cell(table, 11, 6, 6, "标准差")
    _analysis_cell(table, 11, 7, 10, f"{float(stats.get('standard_deviation') or 0):.2f}")
    _analysis_cell(table, 12, 1, 1, "最高分")
    _analysis_cell(table, 12, 2, 3, _score_text(stats.get("maximum")))
    _analysis_cell(table, 12, 4, 5, "最低分")
    _analysis_cell(table, 12, 6, 6, _score_text(stats.get("minimum")))
    _analysis_cell(table, 12, 7, 8, "及格率")
    _analysis_cell(table, 12, 9, 10, f"{float(stats.get('pass_rate') or 0):.2f}%")

    _analysis_cell(table, 13, 0, 1, "阅卷形式(打√)")
    for (start, end), label in zip(row_specs, ("本人阅卷", "同行阅卷", "集体阅卷", "机器阅卷", "其他")):
        suffix = " √" if str(fields.get("marking_form") or "").strip() == label else ""
        _analysis_cell(table, 13, start, end, f"{label}{suffix}")

    _analysis_cell(table, 14, 0, 10, "学生成绩分布图", size=9)
    chart_cell = _analysis_cell(table, 15, 0, 10, "")
    chart_paragraph = chart_cell.paragraphs[0]
    chart_paragraph.alignment = 1
    chart_paragraph.paragraph_format.space_before = 0
    chart_paragraph.paragraph_format.space_after = 0
    chart_paragraph.add_run().add_picture(_chart_image(distribution), width=Cm(16.54), height=Cm(5.16))

    _analysis_cell(table, 16, 0, 0, "")
    _analysis_cell(
        table,
        16,
        1,
        10,
        "简要分析试题结构，成绩分布，学生掌握情况及其主要原因，提出教学改进意见与措施",
        size=9,
        align=0,
    )
    analysis_label = _analysis_cell(table, 17, 0, 0, "试卷分析", size=9, bold=True)
    _set_cell_direction_vertical(analysis_label)
    analysis_cell = _analysis_cell(table, 17, 1, 10, "", align=0)
    _set_analysis_body(analysis_cell, structured.get("analysis_text") or fields.get("analysis_text") or "")

    _analysis_cell(table, 18, 0, 5, "系（教研室）审核意见：", size=9, align=0)
    _analysis_cell(table, 18, 6, 10, "教学院长审核意见：", size=9, align=0)
    _analysis_cell(table, 19, 0, 5, fields.get("department_review_opinion") or "", size=9, align=0)
    _analysis_cell(table, 19, 6, 10, fields.get("dean_review_opinion") or "", size=9, align=0)
    department_signature_cell = _analysis_cell(table, 20, 0, 5, "")
    dean_signature_cell = _analysis_cell(table, 20, 6, 10, "")
    department_ids = fields.get("department_signature_ids") if isinstance(fields.get("department_signature_ids"), list) else []
    dean_ids = fields.get("dean_signature_ids") if isinstance(fields.get("dean_signature_ids"), list) else []
    _set_review_signature(
        department_signature_cell,
        fields.get("department_signature_image_path"),
        len(department_ids) or (1 if fields.get("department_signature_image_path") else 0),
    )
    _set_review_signature(
        dean_signature_cell,
        fields.get("dean_signature_image_path"),
        len(dean_ids) or (1 if fields.get("dean_signature_image_path") else 0),
    )
    _analysis_cell(table, 21, 0, 10, "注：1、本表一式两份，一份交学生所在学院，一份交开课学院存档。", size=9, bold=True, align=0)
    return table


def build_exam_analysis_docx(parse_payload: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    payload = _payload(parse_payload)
    fields = payload.get("fields") if isinstance(payload.get("fields"), dict) else {}
    structured = payload.get("structured") if isinstance(payload.get("structured"), dict) else {}
    document = Document()
    section = document.sections[0]
    section.page_width = Pt(595.25)
    section.page_height = Pt(841.85)
    section.top_margin = Pt(19.4)
    section.bottom_margin = Pt(19.4)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    section.header_distance = Pt(36)
    section.footer_distance = Pt(36)
    document.styles["Normal"].paragraph_format.space_after = Pt(0)
    _build_analysis_table(document, fields, structured)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_academic_final_material_docx(parse_payload: dict[str, Any], template_key: str) -> bytes:
    if template_key == ACADEMIC_GRADE_REGISTER_TYPE:
        return build_grade_register_docx(parse_payload)
    if template_key == ACADEMIC_EXAM_ANALYSIS_TYPE:
        return build_exam_analysis_docx(parse_payload)
    raise ValueError("不支持的教务期末材料类型。")
