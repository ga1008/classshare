"""Render lesson plans to DOCX using the GXUFL printed 教案 layout.

The export intentionally follows the school Word template instead of a generic
report style: cover page with the school header image and underlined fields,
then one fixed-width lesson-plan table per session. The table dimensions,
margins, borders, and font sizes are taken from the reference teaching file so
Word/PDF/print output stays stable.
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from .libreoffice_service import convert_docx_bytes_to_pdf as _safe_docx_to_pdf
from . import lesson_plan_markdown as md

_CN_FONT = "宋体"
_TITLE_FONT = "隶书"
_LATIN_FONT = "Times New Roman"
_MONO_FONT = "Consolas"
_HEADER_IMAGE_PATH = Path(__file__).with_name("assets") / "gxufl_lesson_plan_header.png"

# Reference DOCX section metrics, in English Metric Units.
_PAGE_WIDTH_EMU = 7_560_310
_PAGE_HEIGHT_EMU = 10_692_130
_MARGIN_TOP_EMU = 914_400
_MARGIN_BOTTOM_EMU = 810_260
_MARGIN_LEFT_EMU = 810_260
_MARGIN_RIGHT_EMU = 808_990
_HEADER_DISTANCE_EMU = 540_385
_FOOTER_DISTANCE_EMU = 629_920

# Reference table grids, in twips (dxa).
_OUTER_GRID = (1911, 324, 6095, 1678)
_OUTER_WIDTH = sum(_OUTER_GRID)
_NESTED_ACTIVITY_GRID = (1141, 2684, 2439, 1844)
_NESTED_ACTIVITY_WIDTH = sum(_NESTED_ACTIVITY_GRID)
_HEADER_IMAGE_WIDTH_EMU = 4_909_185
_HEADER_IMAGE_HEIGHT_EMU = 885_825


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------
def _set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    cn: str = _CN_FONT,
    latin: str = _LATIN_FONT,
    color: str | None = None,
    underline: bool = False,
) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = latin
    run.font.underline = underline
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cn)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _write_text_runs(
    paragraph,
    text: Any,
    *,
    size: float = 12,
    bold: bool = False,
    cn: str = _CN_FONT,
    latin: str = _LATIN_FONT,
    color: str | None = None,
    underline: bool = False,
) -> None:
    run = paragraph.add_run(str(text or ""))
    _set_run_font(run, size=size, bold=bold, cn=cn, latin=latin, color=color, underline=underline)


def _set_paragraph_line(paragraph, *, line_twips: int = 400, rule: str = "exact") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), rule)


def _set_first_line_indent(paragraph, twips: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLine"), str(twips))
    ind.set(qn("w:firstLineChars"), "708")


def _set_table_borders(table) -> None:
    tbl_pr = table._element.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _set_table_width_twips(table, width: int) -> None:
    tbl_pr = table._element.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")


def _set_table_fixed_layout(table) -> None:
    tbl_pr = table._element.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_indent_twips(table, indent: int) -> None:
    tbl_pr = table._element.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def _set_table_grid(table, widths: tuple[int, ...]) -> None:
    tbl = table._tbl
    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)


def _configure_table(table, *, width: int, grid: tuple[int, ...]) -> None:
    table.autofit = False
    _set_table_width_twips(table, width)
    _set_table_fixed_layout(table)
    _set_table_grid(table, grid)
    _set_table_borders(table)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(grid):
                _set_cell_width_twips(cell, grid[index])


def _set_cell_width_twips(cell, width: int, *, grid_span: int | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    if grid_span and grid_span > 1:
        span = tc_pr.find(qn("w:gridSpan"))
        if span is None:
            span = OxmlElement("w:gridSpan")
            tc_pr.append(span)
        span.set(qn("w:val"), str(grid_span))


def _set_row_height_twips(row, height: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height))


def _clear_cell(cell) -> None:
    cell.text = ""


def _cell_paragraph(cell, *, first: bool = True):
    if first and cell.paragraphs:
        para = cell.paragraphs[0]
        para.text = ""
        return para
    return cell.add_paragraph()


def _write_runs(
    paragraph,
    text: Any,
    *,
    size: float = 12,
    bold: bool = False,
    cn: str = _CN_FONT,
    latin: str = _LATIN_FONT,
    color: str | None = None,
) -> None:
    for run_spec in md.inline_runs(text):
        run = paragraph.add_run(run_spec["text"])
        _set_run_font(
            run,
            size=size,
            bold=bold or run_spec["bold"],
            cn=cn,
            latin=_MONO_FONT if run_spec.get("code") else latin,
            color=color,
        )


def _write_runs_with_bold_labels(
    paragraph,
    text: Any,
    *,
    size: float = 12,
    bold: bool = False,
    bold_labels: tuple[str, ...] = (),
) -> None:
    raw = str(text or "")
    for label in bold_labels:
        if raw.startswith(label):
            _write_text_runs(paragraph, label, size=size, bold=True)
            remainder = raw[len(label) :]
            if remainder:
                _write_runs(paragraph, remainder, size=size, bold=bold)
            return
    _write_runs(paragraph, raw, size=size, bold=bold)


def _write_plain(
    cell,
    text: Any,
    *,
    size: float = 12,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    vertical=WD_CELL_VERTICAL_ALIGNMENT.TOP,
    line_twips: int = 400,
    bold_labels: tuple[str, ...] = (),
) -> None:
    _clear_cell(cell)
    cell.vertical_alignment = vertical
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines:
        lines = [""]
    for idx, line in enumerate(lines):
        para = _cell_paragraph(cell, first=(idx == 0))
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _set_paragraph_line(para, line_twips=line_twips)
        _write_runs_with_bold_labels(
            para,
            line,
            size=size,
            bold=bold,
            bold_labels=bold_labels,
        )


def _write_label(cell, text: str) -> None:
    _write_plain(
        cell,
        text,
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        vertical=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        line_twips=400,
    )


# ---------------------------------------------------------------------------
# Markdown -> DOCX inside 教学内容及过程
# ---------------------------------------------------------------------------
def _render_markdown_into_cell(cell, markdown: Any) -> None:
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    blocks = md.parse_blocks(markdown)
    first = True
    if not blocks:
        para = _cell_paragraph(cell, first=True)
        _set_paragraph_line(para)
        run = para.add_run("（待补充）")
        _set_run_font(run, size=12)
        return

    for block in blocks:
        btype = block.get("type")
        if btype == "heading":
            para = _cell_paragraph(cell, first=first)
            para.paragraph_format.space_before = Pt(0 if first else 3)
            para.paragraph_format.space_after = Pt(0)
            _set_paragraph_line(para)
            run = para.add_run(str(block.get("text") or ""))
            _set_run_font(run, size=12, bold=True)
        elif btype == "para":
            para = _cell_paragraph(cell, first=first)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            _set_paragraph_line(para)
            _write_runs(para, block.get("text", ""), size=12)
        elif btype in ("ul", "ol"):
            for i, item in enumerate(block.get("items", []), start=1):
                para = _cell_paragraph(cell, first=(first and i == 1))
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                _set_paragraph_line(para)
                prefix = f"{i}. " if btype == "ol" else "•  "
                run = para.add_run(prefix)
                _set_run_font(run, size=12)
                _write_runs(para, item, size=12)
        elif btype == "table":
            _render_md_table_into_cell(cell, block, first=first)
        elif btype == "mermaid":
            _render_mermaid_into_cell(cell, block.get("text", ""), first=first)
        elif btype == "code":
            para = _cell_paragraph(cell, first=first)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            _set_paragraph_line(para)
            run = para.add_run(block.get("text", ""))
            _set_run_font(run, size=10.5, cn=_CN_FONT, latin=_MONO_FONT)
        first = False


def _render_md_table_into_cell(cell, block: dict[str, Any], *, first: bool) -> None:
    header = block.get("header", [])
    rows = block.get("rows", [])
    col_count = max(len(header), max((len(r) for r in rows), default=0)) or 1
    if not first:
        spacer = cell.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        _set_paragraph_line(spacer, line_twips=120)

    nested = cell.add_table(rows=1, cols=col_count)
    nested.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_count == 4:
        grid = _NESTED_ACTIVITY_GRID
    else:
        per_col = max(900, _NESTED_ACTIVITY_WIDTH // col_count)
        grid = tuple(per_col for _ in range(col_count))
    _configure_table(nested, width=sum(grid), grid=grid)

    head_cells = nested.rows[0].cells
    for idx in range(col_count):
        text = header[idx] if idx < len(header) else ""
        _set_cell_width_twips(head_cells[idx], grid[idx])
        _write_plain(head_cells[idx], text, size=12, bold=True, line_twips=400)

    for row in rows:
        body_cells = nested.add_row().cells
        for idx in range(col_count):
            text = row[idx] if idx < len(row) else ""
            _set_cell_width_twips(body_cells[idx], grid[idx])
            _write_plain(
                body_cells[idx],
                str(text).replace("<br>", "\n"),
                size=12,
                bold=False,
                line_twips=400,
            )


def _clean_mermaid_node(raw: str) -> tuple[str, str]:
    text = raw.strip().strip(";")
    match = re.match(r'^([A-Za-z0-9_:-]+)\s*(?:\[\s*"?(.+?)"?\s*\]|\(\s*"?(.+?)"?\s*\)|\{\s*"?(.+?)"?\s*\})?$', text)
    if not match:
        return text, text
    node_id = match.group(1)
    label = next((g for g in match.groups()[1:] if g), node_id)
    return node_id, label.strip().strip('"')


def _parse_mermaid_flowchart(code: Any) -> tuple[str, dict[str, str], list[tuple[str, str]]]:
    direction = "TD"
    nodes: dict[str, str] = {}
    edges: list[tuple[str, str]] = []
    for raw_line in str(code or "").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("%%"):
            continue
        header = re.match(r"^(?:flowchart|graph)\s+([A-Za-z]{2})", line, re.IGNORECASE)
        if header:
            direction = header.group(1).upper()
            continue
        if line.startswith(("style ", "classDef ", "class ")):
            continue
        edge_match = re.match(r"(.+?)\s*(?:--[^-]*-->|-->|---|==>|-.->)\s*(.+)", line)
        if not edge_match:
            node_id, label = _clean_mermaid_node(line)
            if node_id:
                nodes.setdefault(node_id, label)
            continue
        left_id, left_label = _clean_mermaid_node(edge_match.group(1))
        right_id, right_label = _clean_mermaid_node(edge_match.group(2))
        if left_id and right_id:
            nodes.setdefault(left_id, left_label)
            nodes.setdefault(right_id, right_label)
            edges.append((left_id, right_id))
    return direction, nodes, edges


def _diagram_font(size: int):
    from PIL import ImageFont

    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size)
            except Exception:
                continue
    return ImageFont.load_default()


def _wrap_diagram_label(draw, text: str, font, max_width: int) -> list[str]:
    chars = list(str(text or ""))
    lines: list[str] = []
    current = ""
    for ch in chars:
        trial = current + ch
        if current and draw.textbbox((0, 0), trial, font=font)[2] > max_width:
            lines.append(current)
            current = ch
        else:
            current = trial
    if current:
        lines.append(current)
    return lines or [""]


def _render_mermaid_png(code: Any) -> bytes | None:
    """Render common Mermaid flowcharts to PNG for Word-safe embedding."""
    from PIL import Image, ImageDraw

    direction, nodes, edges = _parse_mermaid_flowchart(code)
    if not nodes:
        return None

    incoming = {node: 0 for node in nodes}
    for _, dst in edges:
        incoming[dst] = incoming.get(dst, 0) + 1
    layer: dict[str, int] = {node: 0 for node, count in incoming.items() if count == 0}
    if not layer and nodes:
        first = next(iter(nodes))
        layer[first] = 0
    changed = True
    guard = 0
    while changed and guard < max(1, len(nodes) * len(nodes)):
        guard += 1
        changed = False
        for src, dst in edges:
            candidate = layer.get(src, 0) + 1
            if candidate > layer.get(dst, -1):
                layer[dst] = candidate
                changed = True
    for node in nodes:
        layer.setdefault(node, 0)

    layer_nodes: dict[int, list[str]] = {}
    for node, idx in layer.items():
        layer_nodes.setdefault(idx, []).append(node)
    for group in layer_nodes.values():
        group.sort(key=lambda item: list(nodes).index(item))

    font = _diagram_font(24)
    small_font = _diagram_font(18)
    node_w = 210
    node_h = 76
    gap_x = 80
    gap_y = 58
    layers = sorted(layer_nodes)
    if direction.startswith("LR"):
        width = 80 + len(layers) * node_w + max(0, len(layers) - 1) * gap_x + 80
        height = 80 + max(len(layer_nodes[i]) for i in layers) * node_h + 80
        height += max(0, max(len(layer_nodes[i]) for i in layers) - 1) * gap_y
    else:
        width = 80 + max(len(layer_nodes[i]) for i in layers) * node_w + 80
        width += max(0, max(len(layer_nodes[i]) for i in layers) - 1) * gap_x
        height = 80 + len(layers) * node_h + max(0, len(layers) - 1) * gap_y + 80
    image = Image.new("RGB", (max(width, 420), max(height, 220)), "white")
    draw = ImageDraw.Draw(image)
    positions: dict[str, tuple[int, int, int, int]] = {}

    for layer_idx in layers:
        group = layer_nodes[layer_idx]
        if direction.startswith("LR"):
            x = 80 + layer_idx * (node_w + gap_x)
            total_h = len(group) * node_h + max(0, len(group) - 1) * gap_y
            y0 = (image.height - total_h) // 2
            for offset, node in enumerate(group):
                y = y0 + offset * (node_h + gap_y)
                positions[node] = (x, y, x + node_w, y + node_h)
        else:
            y = 80 + layer_idx * (node_h + gap_y)
            total_w = len(group) * node_w + max(0, len(group) - 1) * gap_x
            x0 = (image.width - total_w) // 2
            for offset, node in enumerate(group):
                x = x0 + offset * (node_w + gap_x)
                positions[node] = (x, y, x + node_w, y + node_h)

    def center(rect: tuple[int, int, int, int]) -> tuple[int, int]:
        return ((rect[0] + rect[2]) // 2, (rect[1] + rect[3]) // 2)

    for src, dst in edges:
        if src not in positions or dst not in positions:
            continue
        sx, sy = center(positions[src])
        dx, dy = center(positions[dst])
        if direction.startswith("LR"):
            start = (positions[src][2], sy)
            end = (positions[dst][0], dy)
            arrow = [(end[0], end[1]), (end[0] - 10, end[1] - 6), (end[0] - 10, end[1] + 6)]
        else:
            start = (sx, positions[src][3])
            end = (dx, positions[dst][1])
            arrow = [(end[0], end[1]), (end[0] - 6, end[1] - 10), (end[0] + 6, end[1] - 10)]
        draw.line([start, end], fill="#4b5563", width=3)
        draw.polygon(arrow, fill="#4b5563")

    for node, rect in positions.items():
        draw.rounded_rectangle(rect, radius=14, fill="#f8fafc", outline="#111827", width=2)
        wrapped = _wrap_diagram_label(draw, nodes.get(node, node), font, node_w - 28)
        line_height = 28
        total_text_h = len(wrapped) * line_height
        y = rect[1] + (node_h - total_text_h) // 2
        for line in wrapped[:2]:
            bbox = draw.textbbox((0, 0), line, font=font)
            x = rect[0] + (node_w - (bbox[2] - bbox[0])) // 2
            draw.text((x, y), line, font=font, fill="#111827")
            y += line_height
        if len(wrapped) > 2:
            draw.text((rect[0] + 16, rect[3] - 23), "...", font=small_font, fill="#64748b")

    out = BytesIO()
    image.save(out, format="PNG")
    return out.getvalue()


def _render_mermaid_into_cell(cell, code: Any, *, first: bool) -> None:
    png = _render_mermaid_png(code)
    if not png:
        para = _cell_paragraph(cell, first=first)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _set_paragraph_line(para)
        run = para.add_run(str(code or ""))
        _set_run_font(run, size=10.5, cn=_CN_FONT, latin=_MONO_FONT)
        return
    para = _cell_paragraph(cell, first=first)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4 if not first else 0)
    para.paragraph_format.space_after = Pt(4)
    _set_paragraph_line(para, line_twips=240, rule="auto")
    para.add_run().add_picture(BytesIO(png), width=Emu(5_650_000))


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Emu(_PAGE_WIDTH_EMU)
    section.page_height = Emu(_PAGE_HEIGHT_EMU)
    section.top_margin = Emu(_MARGIN_TOP_EMU)
    section.bottom_margin = Emu(_MARGIN_BOTTOM_EMU)
    section.left_margin = Emu(_MARGIN_LEFT_EMU)
    section.right_margin = Emu(_MARGIN_RIGHT_EMU)
    section.header_distance = Emu(_HEADER_DISTANCE_EMU)
    section.footer_distance = Emu(_FOOTER_DISTANCE_EMU)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = _LATIN_FONT
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), _LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), _CN_FONT)


def _write_first_page_footer(document: Document) -> None:
    footer = document.sections[0].first_page_footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    _set_paragraph_line(para, line_twips=300, rule="auto")
    _write_text_runs(para, "广西外国语学院教务处 印制", size=15)


def _add_blank_paragraph(document: Document, *, line_twips: int = 240) -> None:
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    _set_paragraph_line(para, line_twips=line_twips)


def _normalize_semester_label(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    text = text.replace(" ", "")
    text = text.replace("第一学期", "第 一 学期")
    text = text.replace("第二学期", "第 二 学期")
    text = text.replace("第三学期", "第 三 学期")
    return text + " "


_CN_NUMBERS = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
}


def _term_to_cn(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text in {"一", "二", "三", "四"}:
        return text
    if "一" in text:
        return "一"
    if "二" in text:
        return "二"
    if "三" in text:
        return "三"
    match = re.search(r"([1-4])", text)
    return _CN_NUMBERS.get(int(match.group(1)), "") if match else text


def _parse_semester_parts(value: Any) -> tuple[str, str]:
    text = str(value or "").strip()
    if not text:
        return "", ""
    normalized = text.replace("－", "-").replace("—", "-").replace("–", "-")
    compact = re.sub(r"\s+", "", normalized)
    match = re.search(r"(\d{4})-(\d{4})(?:学年)?(?:第?([一二三四1-4])学期?|Term([1-4]))?", compact, re.IGNORECASE)
    if match:
        year = f"{match.group(1)}—{match.group(2)}"
        term = _term_to_cn(match.group(3) or match.group(4) or "")
        return year, term
    match = re.search(r"(\d{4})-(\d{4})([1-4])$", compact)
    if match:
        return f"{match.group(1)}—{match.group(2)}", _term_to_cn(match.group(3))
    return text, ""


def _write_semester_line(document: Document, value: Any) -> None:
    year, term = _parse_semester_parts(value)
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(15.6)
    if year:
        _write_text_runs(para, year, size=16, underline=True)
    else:
        _write_text_runs(para, "\u00a0" * 10, size=16, underline=True)
    _write_text_runs(para, "学年第 ", size=16)
    if term:
        _write_text_runs(para, term, size=16, underline=True)
    else:
        _write_text_runs(para, "\u00a0", size=16, underline=True)
    _write_text_runs(para, " 学期", size=16)


def _cover_value_run(paragraph, value: Any, *, min_chars: int, pad_mode: str = "center") -> None:
    raw = str(value or "").strip()
    pad = max(0, min_chars - len(raw))
    if pad_mode == "right":
        left = 0
        right = pad
    else:
        left = pad // 2
        right = pad - left
    # Word does not reliably draw underline for trailing regular spaces. NBSP
    # keeps the school-template field line visible at the intended length.
    blank = "\u00a0"
    text = f"{blank * left}{raw}{blank * right}"
    run = paragraph.add_run(text)
    _set_run_font(run, size=16, underline=True)


def _cover_line(
    document: Document,
    label: str,
    value: Any,
    *,
    min_chars: int = 24,
    pad_mode: str = "center",
) -> None:
    para = document.add_paragraph()
    _set_first_line_indent(para, 2266)
    _set_paragraph_line(para, line_twips=600)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    label_run = para.add_run(label)
    _set_run_font(label_run, size=16)
    _cover_value_run(para, value, min_chars=min_chars, pad_mode=pad_mode)


def _cover_continuation_line(
    document: Document,
    value: str,
    *,
    min_chars: int = 30,
    label_width_chars: int = 5,
    pad_mode: str = "right",
) -> None:
    para = document.add_paragraph()
    _set_first_line_indent(para, 2266)
    _set_paragraph_line(para, line_twips=600)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    placeholder = para.add_run("\u3000" * label_width_chars + "\u00a0")
    _set_run_font(placeholder, size=16)
    _cover_value_run(para, value, min_chars=min_chars, pad_mode=pad_mode)


def _display_width(text: str) -> int:
    width = 0
    for char in text:
        width += 2 if ord(char) > 127 else 1
    return width


def _split_display_lines(text: str, *, max_width: int) -> list[str]:
    remaining = str(text or "").strip()
    if not remaining:
        return [""]
    lines: list[str] = []
    while remaining:
        current = ""
        for index, char in enumerate(remaining):
            if current and _display_width(current + char) > max_width:
                break
            current += char
        if current == remaining:
            lines.append(current.strip())
            break
        split_at = len(current)
        soft = max(current.rfind(" "), current.rfind("，"), current.rfind("、"))
        if soft >= 12:
            split_at = soft + 1
        lines.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    return [line for line in lines if line] or [""]


def _split_textbook_line(textbook: Any, publisher: Any) -> list[str]:
    text = str(textbook or "").strip()
    pub = str(publisher or "").strip()
    if pub and pub not in text:
        text = f"{text} {pub}" if text else pub
    return _split_display_lines(text, max_width=36)


def _write_cover(document: Document, cover: dict[str, Any]) -> None:
    _write_first_page_footer(document)
    _add_blank_paragraph(document, line_twips=700)
    header_para = document.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(0)
    if _HEADER_IMAGE_PATH.exists():
        header_para.add_run().add_picture(
            str(_HEADER_IMAGE_PATH),
            width=Emu(_HEADER_IMAGE_WIDTH_EMU),
            height=Emu(_HEADER_IMAGE_HEIGHT_EMU),
        )

    for _ in range(4):
        _add_blank_paragraph(document, line_twips=315)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(13)
    title.add_run()
    for part in ("教", "  ", "案"):
        run = title.add_run(part)
        _set_run_font(run, size=48, bold=True, cn=_TITLE_FONT)

    _write_semester_line(document, cover.get("semester_label"))

    _add_blank_paragraph(document, line_twips=900)
    _cover_line(document, "课程名称：", cover.get("course_name"), min_chars=16)
    _cover_line(document, "课程类别：", cover.get("course_category"), min_chars=16)

    credit_para = document.add_paragraph()
    _set_first_line_indent(credit_para, 2266)
    _set_paragraph_line(credit_para, line_twips=600)
    credit_label = credit_para.add_run("学    分：")
    _set_run_font(credit_label, size=16)
    spacer = credit_para.add_run("\u00a0" * 2)
    _set_run_font(spacer, size=16)
    _cover_value_run(credit_para, cover.get("credits", ""), min_chars=5)
    spacer = credit_para.add_run("\u00a0" * 2)
    _set_run_font(spacer, size=16)
    hours_label = credit_para.add_run("学   时：")
    _set_run_font(hours_label, size=16)
    _cover_value_run(credit_para, cover.get("total_hours", ""), min_chars=5)

    _cover_line(document, "授课教师：", cover.get("teacher_name"), min_chars=16)
    _cover_line(document, "教学单位：", cover.get("teaching_unit"), min_chars=16)
    _cover_line(document, "授课班级：", cover.get("class_name"), min_chars=16)

    textbook_lines = _split_textbook_line(cover.get("textbook"), cover.get("publisher"))
    if textbook_lines:
        _cover_line(document, "使用教材：", textbook_lines[0], min_chars=18, pad_mode="right")
        for line in textbook_lines[1:]:
            _cover_continuation_line(document, line, min_chars=12, label_width_chars=4)


# ---------------------------------------------------------------------------
# Session tables
# ---------------------------------------------------------------------------
def _cn_number(value: Any) -> str:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return str(value or "").strip()
    if number <= 10:
        return _CN_NUMBERS.get(number, str(number))
    if number < 20:
        return "十" + (_CN_NUMBERS.get(number % 10, "") if number % 10 else "")
    tens = number // 10
    ones = number % 10
    return _CN_NUMBERS.get(tens, str(tens)) + "十" + (_CN_NUMBERS.get(ones, "") if ones else "")


def _weekday_cn(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text.startswith("星期"):
        text = text.replace("星期", "").strip()
    if text in {"一", "二", "三", "四", "五", "六", "日", "天"}:
        return "日" if text == "天" else text
    try:
        number = int(text)
    except ValueError:
        return text
    return {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "日"}.get(number, text)


def _schedule_parts(session: dict[str, Any]) -> dict[str, str]:
    schedule = session.get("schedule") or {}
    text = str(schedule.get("text") or "")
    date_text = str(schedule.get("date") or "").strip()
    year = month = day = ""
    date_match = re.search(r"(\d{4})\s*[-年/\.]\s*(\d{1,2})\s*[-月/\.]\s*(\d{1,2})", date_text or text)
    if date_match:
        year, month, day = date_match.group(1), date_match.group(2).zfill(2), date_match.group(3).zfill(2)
    week = schedule.get("week_index") or ""
    if not week:
        week_match = re.search(r"week\s*([一二三四五六七八九十\d]+)", text, re.IGNORECASE)
        if not week_match:
            week_match = re.search(r"(?:第\s*)?([一二三四五六七八九十\d]+)\s*周", text)
        if week_match:
            week = week_match.group(1)
    weekday = schedule.get("weekday") or ""
    if not weekday:
        weekday_match = re.search(r"星期\s*([一二三四五六日天\d])", text)
        if weekday_match:
            weekday = weekday_match.group(1)
    sections = str(schedule.get("sections") or "").strip()
    if not sections:
        section_match = re.search(r"sections?\s*(\d{1,2}\s*(?:[-~－—]\s*\d{1,2})?)", text, re.IGNORECASE)
        if not section_match:
            section_match = re.search(r"(?:第\s*)?(\d{1,2}\s*(?:[-~－—]\s*\d{1,2})?)\s*节", text)
        if section_match:
            sections = section_match.group(1)
    sections = re.sub(r"\s+", "", sections).replace("－", "-").replace("—", "-").replace("~", "-")
    return {
        "year": year,
        "month": month,
        "day": day,
        "week": _cn_number(week),
        "weekday": _weekday_cn(weekday),
        "sections": sections,
        "raw": text,
    }


def _schedule_text(session: dict[str, Any]) -> str:
    schedule = session.get("schedule") or {}
    if schedule.get("text"):
        return str(schedule["text"])
    parts = [str(schedule.get("date") or "")]
    if schedule.get("week_index"):
        parts.append(f"第 {schedule['week_index']} 周")
    if schedule.get("weekday"):
        parts.append(str(schedule["weekday"]))
    if schedule.get("sections"):
        parts.append(f"第 {schedule['sections']} 节")
    return " ".join(p for p in parts if p)


def _write_schedule_cell(cell, session: dict[str, Any]) -> None:
    parts = _schedule_parts(session)
    if not (parts["year"] and parts["month"] and parts["day"] and parts["week"] and parts["sections"]):
        _write_plain(
            cell,
            _schedule_text(session),
            size=12,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            vertical=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            line_twips=400,
        )
        return
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = _cell_paragraph(cell, first=True)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    _set_paragraph_line(para, line_twips=400)
    _write_text_runs(para, f"{parts['year']} 年 ", size=12)
    _write_text_runs(para, parts["month"], size=12, underline=True)
    _write_text_runs(para, " 月 ", size=12)
    _write_text_runs(para, parts["day"], size=12, underline=True)
    _write_text_runs(para, " 日  第 ", size=12)
    _write_text_runs(para, parts["week"], size=12, underline=True)
    _write_text_runs(para, " 周  星期 ", size=12)
    if parts["weekday"]:
        _write_text_runs(para, parts["weekday"], size=12, underline=True)
    else:
        _write_text_runs(para, "\u00a0", size=12, underline=True)
    _write_text_runs(para, "  第 ", size=12)
    _write_text_runs(para, parts["sections"], size=12, underline=True)
    _write_text_runs(para, " 节", size=12)


def _combine_key_difficult(session: dict[str, Any]) -> str:
    parts = []
    if session.get("key_points"):
        key = str(session["key_points"]).strip()
        parts.append(key if key.startswith("重点") else f"重点： {key}")
    if session.get("difficulties"):
        difficult = str(session["difficulties"]).strip()
        parts.append(difficult if difficult.startswith("难点") else f"难点： {difficult}")
    return "\n".join(parts)


def _combine_methods(session: dict[str, Any]) -> str:
    parts = []
    if session.get("methods"):
        parts.append(str(session["methods"]).strip())
    if session.get("means"):
        parts.append(str(session["means"]).strip())
    return "\n".join(p for p in parts if p)


def _merge_width(cells, start: int, end: int):
    merged = cells[start]
    for idx in range(start + 1, end + 1):
        merged = merged.merge(cells[idx])
    width = sum(_OUTER_GRID[start : end + 1])
    _set_cell_width_twips(merged, width, grid_span=end - start + 1)
    return merged


def _write_label_content_row(
    row,
    *,
    label: str,
    content: Any,
    height: int,
    markdown: bool = False,
    bold_labels: tuple[str, ...] = (),
) -> None:
    _set_row_height_twips(row, height)
    cells = row.cells
    label_cell = _merge_width(cells, 0, 1)
    content_cell = _merge_width(cells, 2, 3)
    _write_label(label_cell, label)
    if markdown:
        _render_markdown_into_cell(content_cell, content)
    else:
        _write_plain(content_cell, content, size=12, line_twips=400, bold_labels=bold_labels)


def _write_session_table(document: Document, session: dict[str, Any], *, page_break_before: bool = False) -> None:
    table = document.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _configure_table(table, width=_OUTER_WIDTH, grid=_OUTER_GRID)
    _set_table_indent_twips(table, 115)
    rows = table.rows

    _set_row_height_twips(rows[0], 453)
    row0_cells = rows[0].cells
    row0_label = _merge_width(row0_cells, 0, 1)
    row0_content = _merge_width(row0_cells, 2, 3)
    _write_label(row0_label, "授课时间")
    if page_break_before and row0_label.paragraphs:
        row0_label.paragraphs[0].paragraph_format.page_break_before = True
    _write_schedule_cell(row0_content, session)
    _write_label_content_row(rows[1], label="授课章节", content=session.get("chapter", ""), height=453)
    _write_label_content_row(
        rows[2],
        label="教学目的和要求",
        content=session.get("objectives", ""),
        height=460,
        bold_labels=("知识目标：", "知识目标:", "能力目标：", "能力目标:", "素养目标：", "素养目标:"),
    )
    _write_label_content_row(
        rows[3],
        label="教学重点和难点",
        content=_combine_key_difficult(session),
        height=450,
        bold_labels=("重点：", "重点:", "难点：", "难点:"),
    )
    _write_label_content_row(rows[4], label="教学方法和手段", content=_combine_methods(session), height=450)

    _set_row_height_twips(rows[5], 631)
    header_cells = rows[5].cells
    process_header = _merge_width(header_cells, 0, 2)
    _set_cell_width_twips(header_cells[3], _OUTER_GRID[3])
    _write_label(process_header, "教学内容及过程")
    _write_label(header_cells[3], "旁批")

    _set_row_height_twips(rows[6], 1134)
    body_cells = rows[6].cells
    process_cell = _merge_width(body_cells, 0, 2)
    _set_cell_width_twips(body_cells[3], _OUTER_GRID[3])
    _render_markdown_into_cell(process_cell, session.get("process", ""))
    _write_plain(body_cells[3], session.get("side_notes", ""), size=12, line_twips=400)

    _write_label_content_row(rows[7], label="教学后记", content=session.get("post_notes", ""), height=390)


def build_lesson_plan_docx(plan: dict[str, Any]) -> bytes:
    """Build DOCX bytes for a hydrated lesson-plan dict."""
    cover = plan.get("cover") or {}
    sessions = plan.get("sessions") or []
    document = Document()
    _setup_document(document)
    _write_cover(document, cover)
    for session in sessions:
        _write_session_table(document, session, page_break_before=True)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF / PNG previews (LibreOffice)
# ---------------------------------------------------------------------------
def convert_docx_to_pdf(docx_bytes: bytes, *, base_name: str = "lesson_plan") -> bytes:
    return _safe_docx_to_pdf(docx_bytes, timeout=120)


def convert_docx_to_png(docx_bytes: bytes, *, base_name: str = "lesson_plan", max_pages: int = 4) -> bytes:
    """Render the first few pages to a single stacked PNG preview."""
    pdf_bytes = convert_docx_to_pdf(docx_bytes, base_name=base_name)
    try:
        import fitz  # PyMuPDF
        from PIL import Image
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError(f"缺少 PNG 渲染依赖：{exc}") from exc

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        images = []
        for page_index in range(min(max_pages, doc.page_count)):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6))
            images.append(Image.frombytes("RGB", (pix.width, pix.height), pix.samples))
    finally:
        doc.close()
    if not images:
        raise RuntimeError("PDF 没有可渲染的页面。")
    width = max(img.width for img in images)
    gap = 16
    height = sum(img.height for img in images) + gap * (len(images) - 1)
    canvas = Image.new("RGB", (width, height), "white")
    y = 0
    for img in images:
        canvas.paste(img, (0, y))
        y += img.height + gap
    out = BytesIO()
    canvas.save(out, format="PNG")
    return out.getvalue()
