import io
import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from .material_final_document_service import (
    ASSESSMENT_PLAN_NOTES,
    FINAL_MATERIAL_TYPES,
    SCORING_RUBRIC_NOTES,
    final_material_label,
    normalize_final_material_payload,
)


@dataclass(frozen=True)
class MaterialExportArtifact:
    content: bytes
    filename: str
    media_type: str


DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
PDF_MEDIA_TYPE = "application/pdf"


TEMPLATE_CONFIGS: dict[str, dict[str, str]] = {
    "teaching_document": {"title": "教学文档", "preferred_format": "docx"},
    "teaching_summary": {"title": "教师教学工作总结", "preferred_format": "docx"},
    "lesson_plan": {"title": "教案", "preferred_format": "docx"},
    "teaching_calendar": {"title": "教学日历", "preferred_format": "xlsx"},
    "evaluation_sheet": {"title": "评学表", "preferred_format": "docx"},
    "final_syllabus": {"title": "课程教学大纲", "preferred_format": "docx"},
    "assessment_plan": {"title": "课程考核计划表", "preferred_format": "docx"},
    "grading_rubric": {"title": "课程考核评分细则", "preferred_format": "docx"},
    "exam_paper": {"title": "课程考核试卷", "preferred_format": "docx"},
    "final_teaching_summary": {"title": "教师教学工作总结", "preferred_format": "docx"},
}

FIELD_LABELS = {
    "school": "学校",
    "college": "学院",
    "department": "系部",
    "course_name": "课程名称",
    "course_code": "课程代码",
    "class_name": "授课班级",
    "teacher_name": "授课教师",
    "academic_year": "学年",
    "semester": "学期",
    "date": "日期",
    "title": "标题",
}


def build_material_export_artifact(
    parse_payload: dict[str, Any],
    *,
    fallback_filename: str,
    requested_format: str | None = None,
) -> MaterialExportArtifact:
    payload = _coerce_payload(parse_payload)
    if "export_payload" not in payload and str(payload.get("template_key") or "").strip():
        template_key = str(payload.get("template_key") or "").strip()
        payload = {
            "document_group": payload.get("document_group") or ("final_material" if template_key in FINAL_MATERIAL_TYPES else ""),
            "document_type": payload.get("document_type") or template_key,
            "document_type_label": payload.get("document_type_label") or final_material_label(template_key),
            "metadata": payload.get("fields") if isinstance(payload.get("fields"), dict) else {},
            "content_markdown": payload.get("content_markdown") or _as_dict(payload.get("structured")).get("rubric_body_markdown") or "",
            "tables": payload.get("tables") if isinstance(payload.get("tables"), list) else [],
            "export_payload": payload,
        }
    export_payload = _as_dict(payload.get("export_payload"))
    template_key = str(export_payload.get("template_key") or payload.get("document_type") or "teaching_document")
    config = TEMPLATE_CONFIGS.get(template_key, TEMPLATE_CONFIGS.get(str(payload.get("document_type")), {}))
    output_format = (requested_format or config.get("preferred_format") or "docx").strip().lower()
    if output_format not in {"docx", "xlsx", "pdf"}:
        output_format = "docx"

    title = _resolve_title(payload, config, fallback_filename)
    base_name = _safe_filename(title or fallback_filename or "材料导出")
    if output_format == "pdf":
        docx_content = _build_docx_export(payload, title=title)
        return MaterialExportArtifact(
            content=_convert_docx_bytes_to_pdf(docx_content, base_name=base_name),
            filename=f"{base_name}.pdf",
            media_type=PDF_MEDIA_TYPE,
        )
    if output_format == "xlsx":
        return MaterialExportArtifact(
            content=_build_xlsx_export(payload, title=title),
            filename=f"{base_name}.xlsx",
            media_type=XLSX_MEDIA_TYPE,
        )
    return MaterialExportArtifact(
        content=_build_docx_export(payload, title=title),
        filename=f"{base_name}.docx",
        media_type=DOCX_MEDIA_TYPE,
    )


def _convert_docx_bytes_to_pdf(docx_content: bytes, *, base_name: str) -> bytes:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("当前服务器未安装 LibreOffice，无法导出 PDF；请先导出 Word 或安装 LibreOffice。")
    with tempfile.TemporaryDirectory(prefix="lanshare-material-export-") as temp_dir:
        work_dir = Path(temp_dir)
        docx_path = work_dir / f"{base_name or 'material'}.docx"
        docx_path.write_bytes(docx_content)
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(work_dir),
            str(docx_path),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=90)
        if result.returncode != 0:
            stderr = (result.stderr or result.stdout or "").strip()
            raise RuntimeError(f"LibreOffice PDF 转换失败：{stderr[:240] or '未知错误'}")
        pdf_path = docx_path.with_suffix(".pdf")
        if not pdf_path.exists():
            pdf_files = sorted(work_dir.glob("*.pdf"))
            pdf_path = pdf_files[0] if pdf_files else pdf_path
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice PDF 转换未生成文件。")
        return pdf_path.read_bytes()


def _build_docx_export(payload: dict[str, Any], *, title: str) -> bytes:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc

    template_key = str(_as_dict(payload.get("export_payload")).get("template_key") or payload.get("document_type") or "")
    if template_key in FINAL_MATERIAL_TYPES:
        return _build_final_material_docx_export(payload, title=title, template_key=template_key)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    if str(payload.get("document_type")) in {"teaching_calendar"}:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    heading = document.add_paragraph()
    heading.alignment = 1
    run = heading.add_run(title)
    _set_run(run, size=18, bold=True)
    heading.paragraph_format.space_after = Pt(14)

    metadata = _as_dict(payload.get("metadata"))
    _add_meta_table(document, metadata)

    content = str(payload.get("content_markdown") or "").strip()
    sections = _normalize_sections(payload, content)
    for section_item in sections:
        section_title = str(section_item.get("title") or "").strip()
        section_content = str(section_item.get("content") or "").strip()
        if not section_content:
            continue
        if section_title and section_title != "正文":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            _set_run(paragraph.add_run(section_title), size=13, bold=True)
        _add_markdown_like_blocks(document, section_content)

    for table_payload in _normalize_tables(payload):
        _add_payload_table(document, table_payload)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 1
    footer_run = footer.add_run(f"由 LanShare 根据解析内容生成 · {datetime.now().strftime('%Y-%m-%d')}")
    _set_run(footer_run, size=8.5, color="64748B")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_final_material_docx_export(payload: dict[str, Any], *, title: str, template_key: str) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc

    export_payload = normalize_final_material_payload(
        document_type=template_key,
        metadata=_as_dict(payload.get("metadata")),
        content_markdown=str(payload.get("content_markdown") or ""),
        tables=_normalize_tables(payload),
        export_payload=_as_dict(payload.get("export_payload")),
    )
    fields = _as_dict(export_payload.get("fields"))
    structured = _as_dict(export_payload.get("structured"))
    layout = _as_dict(export_payload.get("layout_profile"))
    margins = _as_dict(layout.get("margins_cm"))

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(float(margins.get("top") or 1.5))
    section.bottom_margin = Cm(float(margins.get("bottom") or 1.5))
    section.left_margin = Cm(float(margins.get("left") or 1.5))
    section.right_margin = Cm(float(margins.get("right") or 1.5))
    section.footer_distance = Cm(float(margins.get("footer") or 1.5))

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if template_key == "assessment_plan":
        _add_assessment_plan_title_block(document, fields)
    elif template_key == "grading_rubric":
        _add_grading_rubric_title_block(document, fields)
    elif template_key == "exam_paper":
        _add_exam_paper_title_block(document, fields)
    else:
        _add_final_title_block(document, template_key, fields)
    if template_key == "assessment_plan":
        _add_assessment_plan_export_body(document, fields, structured)
    elif template_key == "grading_rubric":
        _add_grading_rubric_export_body(document, fields, structured)
    else:
        _add_exam_paper_export_body(document, fields, structured)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if template_key in {"assessment_plan", "grading_rubric"}:
        _add_docx_field(footer, "PAGE", size=9, color="000000")
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    if template_key == "exam_paper":
        _add_exam_paper_footer(footer)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    footer_label = final_material_label(template_key)
    _set_run_songti(footer.add_run(f"广西外国语学院{footer_label}       第 "), 9, color="808080")
    _add_docx_field(footer, "PAGE", size=9, color="808080")
    _set_run_songti(footer.add_run(" 页 共 "), 9, color="808080")
    _add_docx_field(footer, "NUMPAGES", size=9, color="808080")
    _set_run_songti(footer.add_run(" 页"), 9, color="808080")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_final_title_block(document: Any, template_key: str, fields: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    label = final_material_label(template_key)
    title = f"广西外国语学院{label}"
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    _set_run_songti(paragraph.add_run(title), 18, bold=True)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_songti(period.add_run(_format_period_line(fields)), 14, bold=True)

    if template_key in {"assessment_plan", "grading_rubric"}:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_songti(subtitle.add_run("（非笔试考核）"), 12)
    elif template_key == "exam_paper":
        flags = document.add_paragraph()
        flags.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_songti(flags.add_run(_exam_flags_text(fields)), 12, bold=True)


def _add_assessment_plan_title_block(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(34)
    title.paragraph_format.space_after = Pt(12)
    _set_run_songti(title.add_run("广西外国语学院课程考核计划表"), 18, bold=True)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_after = Pt(4)
    _add_assessment_period_runs(period, fields)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    _set_run_songti(subtitle.add_run(f"（{_assessment_mode_label(fields)}）"), 12)


def _add_grading_rubric_title_block(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(34)
    title.paragraph_format.space_after = Pt(12)
    _set_run_songti(title.add_run("广西外国语学院课程考核评分细则"), 18, bold=True)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_after = Pt(4)
    _add_assessment_period_runs(period, fields)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    _set_run_songti(subtitle.add_run(f"（{_assessment_mode_label(fields)}）"), 12)


def _add_exam_paper_title_block(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(34)
    title.paragraph_format.space_after = Pt(12)
    _set_run_songti(title.add_run("广西外国语学院课程考核试卷"), 18, bold=True)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_after = Pt(12)
    _add_assessment_period_runs(period, fields)

    flags = document.add_paragraph()
    flags.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flags.paragraph_format.space_after = Pt(12)
    _set_run_songti(flags.add_run(_exam_flags_text(fields)), 12, bold=True)
    _add_exam_student_line(document)


def _add_assessment_plan_export_body(document: Any, fields: dict[str, Any], structured: dict[str, Any]) -> None:
    from docx.shared import Pt

    _add_plan_meta_table(document, fields)
    gap = document.add_paragraph()
    gap.paragraph_format.space_before = Pt(10)
    gap.paragraph_format.space_after = Pt(0)
    items = structured.get("assessment_items") if isinstance(structured.get("assessment_items"), list) else []
    _add_assessment_items_table(document, items)
    _add_plan_notes(document)


def _add_grading_rubric_export_body(document: Any, fields: dict[str, Any], structured: dict[str, Any]) -> None:
    from docx.shared import Pt

    _add_rubric_meta_table(document, fields)
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(4)
    _set_run_songti(heading.add_run("评分细则"), 10.5, bold=True)
    body_table, cell = _add_rubric_body_table(document)
    rubric_items = structured.get("rubric_items") if isinstance(structured.get("rubric_items"), list) else []
    body_markdown = str(structured.get("rubric_body_markdown") or "").strip()
    if body_markdown:
        _add_rubric_body_text_to_cell(cell, body_markdown)
    elif rubric_items:
        _add_rubric_items_to_cell(cell, rubric_items)
    else:
        for section in structured.get("sections") or []:
            title = str(section.get("title") or "").strip()
            content = str(section.get("content") or "").strip()
            if title:
                _add_rubric_line_to_cell(cell, title, force_bold=True)
            _add_rubric_body_text_to_cell(cell, content)
    _add_rubric_notes(document)


def _add_exam_paper_export_body(document: Any, fields: dict[str, Any], structured: dict[str, Any]) -> None:
    _add_exam_meta_table(document, fields)
    _add_exam_score_summary_table(document, structured)
    paper_sections = structured.get("paper_sections") if isinstance(structured.get("paper_sections"), list) else []
    if not paper_sections:
        paper_sections = structured.get("sections") if isinstance(structured.get("sections"), list) else []
    for index, section in enumerate(paper_sections, start=1):
        title = str(section.get("title") or f"第{index}题").strip()
        score = str(section.get("score") or "").strip()
        heading_text = title if not score or score in title else f"{title}（共 {score} 分）"
        _add_exam_section_heading(document, heading_text)
        _add_exam_content_blocks(document, str(section.get("content") or ""))


def _add_plan_meta_table(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    rows = [
        ["课程名称", _field(fields, "course_name"), "", ""],
        ["专业年级班级", _field(fields, "class_name"), "考核类型", _checked_pair("考查", "考试", _field(fields, "assessment_type") or "考试")],
        ["命题教师", _signature_value(fields, "examiner_name", "teacher_name", signature_key="examiner_signature"), "系（教研室）主任\n审核签字", _signature_value(fields, "reviewer_name", signature_key="reviewer_signature")],
        ["命题日期", _field(fields, "date"), "", ""],
    ]
    table = document.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    widths = [4.45, 4.15, 4.05, 4.85]
    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _cm(1.12)
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            _set_cell_text(cell, value, bold=True, align=1)
            _set_cell_width(cell, widths[col_index])
        if row_index in {0, 3}:
            row.cells[1].merge(row.cells[3])


def _add_rubric_meta_table(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    rows = [
        ["课程名称", _field(fields, "course_name"), "", ""],
        ["专业年级班级", _field(fields, "class_name"), "", ""],
        ["考核形式", _field(fields, "assessment_type") or _field(fields, "assessment_method") or "考试", "命题日期", _field(fields, "date")],
        ["命题教师", _signature_value(fields, "examiner_name", "teacher_name", signature_key="examiner_signature"), "系（教研室）主任\n审核签字", _signature_value(fields, "reviewer_name", signature_key="reviewer_signature")],
    ]
    table = document.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    widths = [3.75, 4.75, 3.65, 5.25]
    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _cm(0.9)
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            _set_cell_text(cell, value, bold=True, align=1)
            _set_cell_width(cell, widths[col_index])
        if row_index in {0, 1}:
            row.cells[1].merge(row.cells[3])


def _add_exam_meta_table(document: Any, fields: dict[str, Any]) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    rows = [
        [(0, 1, "课程名称"), (2, 10, _field(fields, "course_name"))],
        [
            (0, 1, "学历层次"),
            (2, 5, _checked_pair("本科", "专科", _field(fields, "education_level") or "本科")),
            (6, 7, "考核类型"),
            (8, 10, _checked_pair("考查", "考试", _field(fields, "assessment_type") or "考试")),
        ],
        [
            (0, 1, "专业年级班级"),
            (2, 5, _field(fields, "class_name")),
            (6, 7, "考试时间"),
            (8, 10, f"（ {_field(fields, 'exam_duration') or '120'} ）分钟"),
        ],
        [
            (0, 1, "试卷类型"),
            (2, 5, _checked_pair("开卷", "闭卷", _field(fields, "paper_type") or "开卷")),
            (6, 7, "命题教师"),
            (8, 10, _field(fields, "examiner_name", "teacher_name")),
        ],
        [
            (0, 1, "系（教研室）\n主任"),
            (2, 5, _field(fields, "reviewer_name")),
            (6, 9, "二级学院（部）\n主管教学领导"),
            (10, 10, _field(fields, "leader_name")),
        ],
    ]
    grid_widths = [1.76, 0.93, 0.75, 1.75, 1.75, 0.69, 1.06, 1.0, 0.75, 1.75, 2.5]
    table = document.add_table(rows=len(rows), cols=len(grid_widths))
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    for row_index, spans in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _cm(0.73 if row_index < 4 else 0.86)
        for col_index, width in enumerate(grid_widths):
            _set_cell_width(row.cells[col_index], width)
        for start, end, value in spans:
            cell = row.cells[start]
            if end > start:
                cell = cell.merge(row.cells[end])
            _set_cell_width(cell, sum(grid_widths[start:end + 1]))
            _set_cell_text(cell, value, bold=start in {0, 6}, align=1, size=10)
    gap = document.add_paragraph()
    gap.paragraph_format.space_after = _pt(8)


def _add_assessment_items_table(document: Any, items: list[dict[str, Any]]) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    headers = ["考核形式", "考核技能/内容", "分 值"]
    widths = [4.45, 9.65, 3.35]
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = _cm(1.0)
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True, align=1)
        _set_cell_width(table.rows[0].cells[index], widths[index])
    for item in items:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _cm(1.75)
        cells = row.cells
        values = [
            item.get("assessment_form") or item.get("form") or "机试",
            item.get("content") or item.get("assessment_content") or "",
            item.get("score") or "",
        ]
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value, align=1 if index != 1 else 0)
            _set_cell_width(cells[index], widths[index])


def _add_exam_score_summary_table(document: Any, structured: dict[str, Any]) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    sections = structured.get("paper_sections") if isinstance(structured.get("paper_sections"), list) else []
    scores = [str(item.get("score") or "").strip() for item in sections if str(item.get("score") or "").strip()]
    if not scores:
        scores = ["100"]
    scores = scores[:5]
    total = structured.get("total_score") or "100"
    section_labels = [_exam_section_number_label(index) for index in range(1, len(scores) + 1)]
    table = document.add_table(rows=3, cols=8)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    widths = [1.76, 1.68, 1.75, 1.75, 1.75, 1.75, 1.75, 2.5]
    headers = ["题号", *section_labels, *[""] * (5 - len(section_labels)), "总分", "核分人"]
    full = ["满分", *scores, *[""] * (5 - len(scores)), _score_text(total), ""]
    actual = ["实得分", *[""] * 5, "", ""]
    for row_index, row_values in enumerate([headers, full, actual]):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _cm(0.72 if row_index else 0.86)
        for col_index, value in enumerate(row_values):
            _set_cell_width(row.cells[col_index], widths[col_index])
            _set_cell_text(row.cells[col_index], value, bold=row_index != 2, align=1, size=10)
    gap = document.add_paragraph()
    gap.paragraph_format.space_after = _pt(8)


def _add_exam_score_box(document: Any) -> None:
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = 0
    table.autofit = False
    _set_table_borders(table)
    widths = [1.44, 2.25]
    for row in table.rows:
        row.height_rule = 1
        row.height = _cm(0.65)
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
    _set_cell_text(table.rows[0].cells[0], "得分", bold=True, align=1, size=11)
    _set_cell_text(table.rows[0].cells[1], "评卷人", bold=True, align=1, size=11)
    _set_cell_text(table.rows[1].cells[0], "", align=1, size=11)
    _set_cell_text(table.rows[1].cells[1], "", align=1, size=11)


def _add_exam_section_heading(document: Any, heading_text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    layout = document.add_table(rows=2, cols=3)
    layout.alignment = 0
    layout.autofit = False
    _set_table_borders(layout, color="FFFFFF", size=0)
    widths = [1.44, 2.25, 10.64]
    for index, width in enumerate(widths):
        layout.columns[index].width = _cm(width)
    for row in layout.rows:
        row.height_rule = 1
        row.height = _cm(0.65)
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            _set_cell_borders(cell, size=0)
    _set_cell_text(layout.rows[0].cells[0], "得分", bold=True, align=1, size=11)
    _set_cell_text(layout.rows[0].cells[1], "评卷人", bold=True, align=1, size=11)
    _set_cell_text(layout.rows[1].cells[0], "", align=1, size=11)
    _set_cell_text(layout.rows[1].cells[1], "", align=1, size=11)
    for cell in (layout.rows[0].cells[0], layout.rows[0].cells[1], layout.rows[1].cells[0], layout.rows[1].cells[1]):
        _set_cell_borders(cell)
    right = layout.rows[0].cells[2].merge(layout.rows[1].cells[2])
    _set_cell_borders(right, size=0)
    _clear_cell(right)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = _pt(6)
    p.paragraph_format.space_after = _pt(2)
    _set_run_songti(p.add_run(str(heading_text or "试题")), 12, bold=True, font_name="黑体")


def _add_exam_score_box_to_cell(cell: Any) -> None:
    table = cell.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_borders(table)
    widths = [1.44, 2.25]
    for row in table.rows:
        row.height_rule = 1
        row.height = _cm(0.65)
        for index, item_cell in enumerate(row.cells):
            _set_cell_width(item_cell, widths[index])
    _set_cell_text(table.rows[0].cells[0], "得分", bold=True, align=1, size=11)
    _set_cell_text(table.rows[0].cells[1], "评卷人", bold=True, align=1, size=11)
    _set_cell_text(table.rows[1].cells[0], "", align=1, size=11)
    _set_cell_text(table.rows[1].cells[1], "", align=1, size=11)


def _add_exam_content_blocks(document: Any, content: str) -> None:
    lines = str(content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    in_fence = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        _add_exam_body_line(document, text)
        paragraph_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        if code_buffer:
            _add_exam_code_block(document, "\n".join(code_buffer).strip())
        code_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_fence:
                flush_code()
                in_fence = False
            else:
                flush_paragraph()
                in_fence = True
            continue
        if in_fence:
            code_buffer.append(line)
            continue
        if not stripped:
            flush_paragraph()
            continue
        if _looks_like_exam_code_line(stripped):
            flush_paragraph()
            _add_exam_code_block(document, stripped)
            continue
        if _starts_new_exam_line(stripped):
            flush_paragraph()
            _add_exam_body_line(document, stripped)
            continue
        paragraph_buffer.append(stripped)
    flush_paragraph()
    flush_code()


def _add_exam_body_line(document: Any, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = _pt(0)
    p.paragraph_format.space_after = _pt(3)
    p.paragraph_format.line_spacing = 1.18
    stripped = str(text or "").strip()
    if re.match(r"^[0-9]+[\.、．]", stripped):
        p.paragraph_format.left_indent = _cm(0.55)
        p.paragraph_format.first_line_indent = _cm(-0.35)
        bold = False
    elif re.match(r"^[a-zA-Z][\.、．]", stripped):
        p.paragraph_format.left_indent = _cm(0.78)
        p.paragraph_format.first_line_indent = _cm(-0.28)
        bold = False
    elif stripped.startswith("•") or stripped.startswith("·"):
        p.paragraph_format.left_indent = _cm(0.9)
        p.paragraph_format.first_line_indent = _cm(-0.28)
        bold = False
    else:
        p.paragraph_format.left_indent = _cm(0.55)
        bold = bool(re.match(r"^[一二三四五六七八九十]+[、.．]|任务\s*\d+[:：]", stripped))
    _set_run_songti(p.add_run(stripped), 10.5, bold=bold)


def _add_exam_code_block(document: Any, text: str) -> None:
    if not str(text or "").strip():
        return
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table, color="E5E7EB", size=4)
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, 13.9)
    _set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
    _shade_docx_cell(cell, "F5F6F8")
    cell.text = ""
    for raw in str(text).splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_before = _pt(0)
        p.paragraph_format.space_after = _pt(0)
        _set_run_songti(p.add_run(raw), 10, font_name="Consolas")


def _starts_new_exam_line(text: str) -> bool:
    return bool(
        re.match(r"^(?:[0-9]+[\.、．]|[a-zA-Z][\.、．]|[（(][0-9]+[）)]|[•·]|任务\s*\d+[:：]|[一二三四五六七八九十]+[、.．])", str(text or "").strip())
    )


def _looks_like_exam_code_line(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    if re.match(r"^\(?['\"]?[A-Za-z0-9_@./:-]+['\"]?\s*,", stripped):
        return True
    return bool(re.search(r"\b(?:SELECT|INSERT|UPDATE|DELETE|CREATE|ALTER|DROP|JOIN|VALUES|mysql|jdbc:|http://|https://)\b", stripped, flags=re.IGNORECASE))


def _add_exam_student_line(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_songti(p.add_run("姓名：_____________  学号：_____________    年级、专业、班级：_________________     座位号：____________"), 10.5, font_name="黑体")
    seal = document.add_paragraph()
    seal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_songti(seal.add_run("-密----------------封----------------线------------------内------------------不--------------------要----------------------答------------------题----------------"), 9, bold=True)


def _add_plan_notes(document: Any) -> None:
    for line in ASSESSMENT_PLAN_NOTES:
        p = document.add_paragraph()
        p.paragraph_format.space_before = _pt(0)
        p.paragraph_format.space_after = _pt(0)
        p.paragraph_format.line_spacing = 1.08
        _set_run_songti(p.add_run(line), 10.5)


def _add_rubric_notes(document: Any) -> None:
    for line in SCORING_RUBRIC_NOTES:
        p = document.add_paragraph()
        p.paragraph_format.space_before = _pt(0)
        p.paragraph_format.space_after = _pt(0)
        p.paragraph_format.line_spacing = 1.08
        _set_run_songti(p.add_run(line), 10.5)


def _add_text_blocks_to_cell(cell: Any, text: str) -> None:
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        p = cell.add_paragraph()
        _set_run_songti(p.add_run(line), 10.5, bold=bool(re.match(r"^[一二三四五六七八九十]+[、.．]|任务\\d+", line)))


def _split_rubric_paragraphs(text: str) -> list[str]:
    raw = " ".join(str(text or "").split())
    if not raw:
        return []
    raw = re.sub(
        r"\s+(?=(?:文件命名与组织|关键信息一致性|给分原则|拼写错误|结果导向|[一二三四五六七八九十]+、|任务\s*\d+[:：]|[0-9](?:\.[0-9])?\s+|[（(]\d+[）)]|【\d+(?:\.\d+)?分】))",
        "\n",
        raw,
    )
    chunks = [line.strip() for line in raw.splitlines() if line.strip()]
    expanded: list[str] = []
    for chunk in chunks:
        parts = re.split(r"(?=【\d+(?:\.\d+)?分】)", chunk)
        for part in parts:
            cleaned = part.strip()
            if cleaned:
                expanded.append(cleaned)
    return expanded or [raw]


def _add_rubric_body_table(document: Any) -> tuple[Any, Any]:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    _set_table_borders(table)
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, 17.4)
    _set_cell_margins(cell, top=92, bottom=92, left=132, right=132)
    cell.text = ""
    return table, cell


def _add_rubric_body_text_to_cell(cell: Any, text: str) -> None:
    for line in _rubric_body_lines(text):
        _add_rubric_line_to_cell(cell, line)


def _add_rubric_items_to_cell(cell: Any, rubric_items: list[dict[str, Any]]) -> None:
    for item in rubric_items:
        title = str(item.get("title") or "评分项目").strip()
        score = str(item.get("score") or "").strip()
        _add_rubric_line_to_cell(cell, title if not score or score in title else f"{title}（共{score}分）", force_bold=True)
        for criterion in item.get("criteria") or []:
            score_text = str(criterion.get("score") or "").strip()
            text = str(criterion.get("text") or "").strip()
            if not text and not score_text:
                continue
            _add_rubric_line_to_cell(cell, f"【{score_text}分】 {text}" if score_text else text)


def _rubric_body_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw_line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").splitlines():
        cleaned = raw_line.strip()
        if not cleaned:
            lines.append("")
            continue
        cleaned = re.sub(r"^#{1,4}\s*", "", cleaned)
        chunks = _split_rubric_paragraphs(cleaned)
        lines.extend(chunks if chunks else [cleaned])
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return lines


def _add_rubric_line_to_cell(cell: Any, line: str, *, force_bold: bool = False) -> None:
    p = cell.add_paragraph()
    text = str(line or "").strip()
    if not text:
        p.paragraph_format.space_after = _pt(3)
        return
    p.paragraph_format.space_before = _pt(0)
    p.paragraph_format.space_after = _pt(2)
    p.paragraph_format.line_spacing = 1.18
    if re.match(r"^【\d+(?:\.\d+)?分】", text):
        p.paragraph_format.left_indent = _cm(1.05)
    elif re.match(r"^[（(]\d+[）)]", text):
        p.paragraph_format.left_indent = _cm(0.65)
    elif re.match(r"^\d+(?:\.\d+)?[\.、．]", text):
        p.paragraph_format.left_indent = _cm(0.35)
    else:
        p.paragraph_format.left_indent = _cm(0)
    bold = force_bold or bool(
        re.match(r"^[一二三四五六七八九十]+[、.．]|任务\s*\d+[:：]?|.*扣分项与给分原则$", text)
    )
    match = re.match(r"^(【\d+(?:\.\d+)?分】)(.*)$", text)
    if match:
        _set_run_songti(p.add_run(match.group(1)), 10.5, bold=True)
        _set_run_songti(p.add_run(match.group(2)), 10.5, bold=bold)
    else:
        _set_run_songti(p.add_run(text), 10.5, bold=bold)


def _format_period_line(fields: dict[str, Any]) -> str:
    academic_year = str(_field(fields, "academic_year") or "").strip()
    start = "20__"
    end = "20__"
    match = re.search(r"(20\d{2}).*?(20\d{2})", academic_year)
    if match:
        start, end = match.group(1), match.group(2)
    semester = str(_field(fields, "semester") or "").strip()
    if "一" in semester or semester == "1":
        sem = "一"
    elif "二" in semester or semester == "2":
        sem = "二"
    else:
        sem = semester or "__"
    return f"（{start[:2]} {start[2:]}  —  {end[:2]} {end[2:]}  学年度第 {sem} 学期）"


def _add_assessment_period_runs(paragraph: Any, fields: dict[str, Any]) -> None:
    start, end, semester = _assessment_period_parts(fields)
    _set_run_songti(paragraph.add_run("（"), 14, bold=True)
    _set_run_songti(paragraph.add_run(start[:2]), 14, bold=True)
    _set_run_songti(paragraph.add_run(start[2:] or "__"), 14, bold=True, underline=True)
    _set_run_songti(paragraph.add_run("  —  "), 14, bold=True)
    _set_run_songti(paragraph.add_run(end[:2]), 14, bold=True)
    _set_run_songti(paragraph.add_run(end[2:] or "__"), 14, bold=True, underline=True)
    _set_run_songti(paragraph.add_run("  学年度第"), 14, bold=True)
    _set_run_songti(paragraph.add_run(semester or "    "), 14, bold=True, underline=True)
    _set_run_songti(paragraph.add_run("学期）"), 14, bold=True)


def _assessment_period_parts(fields: dict[str, Any]) -> tuple[str, str, str]:
    academic_year = str(_field(fields, "academic_year") or "").strip()
    start = "20__"
    end = "20__"
    match = re.search(r"(20\d{2}).*?(20\d{2})", academic_year)
    if match:
        start, end = match.group(1), match.group(2)
    semester = str(_field(fields, "semester") or "").strip()
    if "一" in semester or semester == "1":
        sem = "一"
    elif "二" in semester or semester == "2":
        sem = "二"
    else:
        sem = semester.replace("第", "").replace("学期", "").strip()
    return start, end, sem


def _assessment_mode_label(fields: dict[str, Any]) -> str:
    raw = str(_field(fields, "assessment_mode_label", "assessment_mode", "academic_exam_mode") or "")
    if "笔试" in raw and "非笔试" not in raw:
        return "笔试考核"
    return "非笔试考核"


def _exam_flags_text(fields: dict[str, Any]) -> str:
    text = str(fields.get("exam_flags") or fields.get("exam_kind") or "期末考试").strip()
    return f"期末考试（ {'√' if '期末' in text else ' '} ）    补考（ {'√' if '补考' in text else ' '} ）    重新学习考试（ {'√' if '重新' in text else ' '} ）"


def _checked_pair(left: str, right: str, selected: str) -> str:
    raw = str(selected or "")
    left_checked = "√" if left in raw and right not in raw else " "
    right_checked = "√" if right in raw or not raw else " "
    return f"{left}（ {left_checked} ）/ {right}（ {right_checked} ）"


def _field(fields: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = fields.get(key)
        if value not in (None, "", [], {}):
            return _stringify(value)
    return ""


def _signature_value(fields: dict[str, Any], *keys: str, signature_key: str) -> str:
    name = _field(fields, *keys)
    signature = _field(fields, signature_key)
    if signature and signature != name:
        return f"{name}    {signature}".strip()
    return name


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, align: int = 0, size: float = 10.5) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = _pt(0)
    paragraph.paragraph_format.space_after = _pt(0)
    _set_run_songti(paragraph.add_run(_stringify(text)), size, bold=bold)
    _set_cell_margins(cell, top=70, bottom=70, left=90, right=90)


def _clear_cell(cell: Any) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()


def _set_cell_width(cell: Any, width_cm: float) -> None:
    cell.width = _cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        from docx.oxml import OxmlElement

        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    from docx.oxml.ns import qn

    tc_w.set(qn("w:w"), str(int(float(width_cm) * 567)))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_borders(table: Any, *, color: str = "000000", size: int = 8) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil" if int(size) <= 0 else "single")
        element.set(qn("w:sz"), str(max(0, int(size))))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_borders(cell: Any, *, color: str = "000000", size: int = 8) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil" if int(size) <= 0 else "single")
        element.set(qn("w:sz"), str(max(0, int(size))))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell: Any, *, top: int = 80, bottom: int = 80, left: int = 100, right: int = 100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run_songti(
    run: Any,
    size: float,
    *,
    bold: bool = False,
    color: str = "000000",
    font_name: str = "宋体",
    underline: bool = False,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)


def _add_docx_field(paragraph: Any, field_code: str, *, size: float = 9, color: str = "808080") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(begin)
    run._element.append(instr)
    run._element.append(end)
    _set_run_songti(run, size, color=color)


def _cm(value: float):
    from docx.shared import Cm

    return Cm(float(value))


def _score_text(value: Any) -> str:
    match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
    if not match:
        return _stringify(value)
    number = float(match.group(0))
    return str(int(number)) if number.is_integer() else str(number)


def _chinese_index(index: int) -> str:
    labels = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 0 <= int(index) < len(labels):
        return labels[int(index)]
    return str(index)


def _exam_section_number_label(index: int) -> str:
    return _chinese_index(index) if int(index) <= 10 else str(index)


def _add_exam_paper_footer(footer: Any) -> None:
    paragraph = footer
    paragraph.alignment = 1
    _set_run_songti(paragraph.add_run("广西外国语学院课程考核试卷          第 "), 9, color="000000")
    _add_docx_field(paragraph, "PAGE", size=9, color="000000")
    _set_run_songti(paragraph.add_run(" 页 共 "), 9, color="000000")
    _add_docx_field(paragraph, "NUMPAGES", size=9, color="000000")
    _set_run_songti(paragraph.add_run(" 页考试过程中不得将试卷拆开"), 9, color="000000")


def _build_xlsx_export(payload: dict[str, Any], *, title: str) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError(f"缺少 XLSX 导出依赖 openpyxl: {exc}") from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "材料内容"
    thin = Side(style="thin", color="D8E0EA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="EAF2FF")
    label_fill = PatternFill("solid", fgColor="F8FAFC")

    ws.merge_cells("A1:F1")
    ws["A1"] = title
    ws["A1"].font = Font(name="Microsoft YaHei", size=16, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    row = 3
    metadata = _as_dict(payload.get("metadata"))
    for key, value in _iter_visible_metadata(metadata):
        ws.cell(row=row, column=1, value=FIELD_LABELS.get(key, key))
        ws.cell(row=row, column=2, value=_stringify(value))
        ws.cell(row=row, column=1).fill = label_fill
        for col in range(1, 3):
            cell = ws.cell(row=row, column=col)
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=True)
        row += 1

    row += 1
    for section in _normalize_sections(payload, str(payload.get("content_markdown") or "")):
        title_cell = ws.cell(row=row, column=1, value=section.get("title") or "正文")
        title_cell.fill = header_fill
        title_cell.font = Font(name="Microsoft YaHei", bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        row += 1
        content_cell = ws.cell(row=row, column=1, value=str(section.get("content") or "").strip())
        content_cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        ws.row_dimensions[row].height = min(220, max(48, len(str(content_cell.value or "")) // 6))
        row += 2

    for table_payload in _normalize_tables(payload):
        rows = _table_rows(table_payload)
        if not rows:
            continue
        title_text = str(table_payload.get("title") or "结构化表格")
        ws.cell(row=row, column=1, value=title_text).font = Font(name="Microsoft YaHei", bold=True)
        row += 1
        for row_values in rows:
            for col_index, value in enumerate(row_values, start=1):
                cell = ws.cell(row=row, column=col_index, value=value)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="center")
                if row_values is rows[0]:
                    cell.fill = header_fill
                    cell.font = Font(name="Microsoft YaHei", bold=True)
            row += 1
        row += 1

    for col in range(1, 9):
        ws.column_dimensions[get_column_letter(col)].width = 18 if col <= 2 else 24
    ws.freeze_panes = "A3"

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def _add_meta_table(document: Any, metadata: dict[str, Any]) -> None:
    visible = list(_iter_visible_metadata(metadata))
    if not visible:
        return
    table = document.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = 1
    for index in range(0, len(visible), 2):
        cells = table.add_row().cells
        pair = visible[index:index + 2]
        for pair_index, (key, value) in enumerate(pair):
            label_cell = cells[pair_index * 2]
            value_cell = cells[pair_index * 2 + 1]
            label_cell.text = FIELD_LABELS.get(key, key)
            value_cell.text = _stringify(value)
            _shade_docx_cell(label_cell, "F8FAFC")
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    _set_run(run, size=9.5, bold=True)
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    _set_run(run, size=9.5)
    document.add_paragraph()


def _add_markdown_like_blocks(document: Any, content: str) -> None:
    lines = str(content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    table_buffer: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph(" ".join(item.strip() for item in paragraph_buffer if item.strip()))
        paragraph.paragraph_format.space_after = _pt(4)
        for run in paragraph.runs:
            _set_run(run, size=10.5)
        paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if len(table_buffer) >= 2:
            _add_payload_table(document, {"title": "", "rows": [_split_markdown_row(line) for line in table_buffer if "|" in line]})
        table_buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            flush_table()
            continue
        if "|" in stripped and re.match(r"^\|?.+\|.+\|?$", stripped):
            flush_paragraph()
            if not re.match(r"^\|?\s*:?-{3,}:?", stripped):
                table_buffer.append(stripped)
            continue
        flush_table()
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            paragraph = document.add_paragraph()
            _set_run(paragraph.add_run(heading.group(2)), size=13 - min(2, len(heading.group(1))), bold=True)
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(bullet.group(1))
            for run in paragraph.runs:
                _set_run(run, size=10.5)
            continue
        paragraph_buffer.append(stripped)
    flush_paragraph()
    flush_table()


def _add_payload_table(document: Any, table_payload: dict[str, Any]) -> None:
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc
    rows = _table_rows(table_payload)
    if not rows:
        return
    title = str(table_payload.get("title") or "").strip()
    if title:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = _pt(6)
        paragraph.paragraph_format.space_after = _pt(3)
        _set_run(paragraph.add_run(title), size=11.5, bold=True)

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_index in range(column_count):
        cell = table.rows[0].cells[col_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_docx_cell(cell, "EAF2FF")
        cell.text = rows[0][col_index] if col_index < len(rows[0]) else ""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _set_run(run, size=9, bold=True)
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for col_index in range(column_count):
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[col_index].text = row_values[col_index] if col_index < len(row_values) else ""
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    _set_run(run, size=8.5)
    document.add_paragraph()


def _normalize_sections(payload: dict[str, Any], content: str) -> list[dict[str, str]]:
    export_payload = _as_dict(payload.get("export_payload"))
    sections = export_payload.get("sections")
    if isinstance(sections, list) and sections:
        normalized = []
        for item in sections:
            if isinstance(item, dict):
                normalized.append({"title": str(item.get("title") or "正文"), "content": str(item.get("content") or "")})
        if normalized:
            return normalized
    return [{"title": "正文", "content": content}]


def _normalize_tables(payload: dict[str, Any]) -> list[dict[str, Any]]:
    tables = payload.get("tables")
    if not isinstance(tables, list):
        tables = _as_dict(payload.get("export_payload")).get("tables")
    if not isinstance(tables, list):
        return []
    return [item for item in tables if isinstance(item, dict)]


def _table_rows(table_payload: dict[str, Any]) -> list[list[str]]:
    rows = table_payload.get("rows")
    if not isinstance(rows, list):
        return []
    normalized = []
    for row in rows:
        if isinstance(row, dict):
            normalized.append([_stringify(value) for value in row.values()])
        elif isinstance(row, list):
            normalized.append([_stringify(value) for value in row])
    max_width = max((len(row) for row in normalized), default=0)
    return [row + [""] * (max_width - len(row)) for row in normalized if any(cell.strip() for cell in row)]


def _iter_visible_metadata(metadata: dict[str, Any]):
    skipped = {"source_filename", "document_group", "document_type"}
    preferred = [
        "school",
        "college",
        "department",
        "course_name",
        "course_code",
        "class_name",
        "teacher_name",
        "academic_year",
        "semester",
        "date",
    ]
    yielded = set()
    for key in preferred:
        value = metadata.get(key)
        if value not in (None, "", [], {}):
            yielded.add(key)
            yield key, value
    for key, value in metadata.items():
        if key in yielded or key in skipped or value in (None, "", [], {}):
            continue
        yield key, value


def _resolve_title(payload: dict[str, Any], config: dict[str, str], fallback_filename: str) -> str:
    metadata = _as_dict(payload.get("metadata"))
    for key in ("title", "course_name"):
        value = str(metadata.get(key) or "").strip()
        if value:
            template_title = config.get("title") or str(payload.get("document_type_label") or "")
            return value if template_title in value else f"{value}-{template_title}" if template_title else value
    return config.get("title") or str(payload.get("document_type_label") or "") or fallback_filename


def _coerce_payload(payload: dict[str, Any] | str | None) -> dict[str, Any]:
    if isinstance(payload, dict):
        return payload
    if isinstance(payload, str) and payload.strip():
        try:
            value = json.loads(payload)
            return value if isinstance(value, dict) else {}
        except json.JSONDecodeError:
            return {}
    return {}


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "-", str(value or "材料导出")).strip(" .")
    return cleaned[:120] or "材料导出"


def _split_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _set_run(run: Any, *, size: float = 10.5, bold: bool = False, color: str = "111827") -> None:
    try:
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade_docx_cell(cell: Any, fill: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _pt(value: float):
    try:
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(f"缺少 DOCX 导出依赖 python-docx: {exc}") from exc
    return Pt(value)
