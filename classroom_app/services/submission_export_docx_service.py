from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Sequence

from fastapi import HTTPException
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .submission_file_alignment import resolve_submission_file_path


DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
BODY_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"
QUESTION_MAX_IMAGE_WIDTH_IN = 5.35
QUESTION_MAX_IMAGE_HEIGHT_IN = 3.35


@dataclass
class ExportAttachment:
    name: str
    relative_path: str
    mime_type: str
    file_ext: str
    path: Path | None
    is_image: bool
    matched_to_question: bool = False


@dataclass
class ExportQuestion:
    section_name: str
    number: int
    question_id: str
    question_type: str
    text: str
    options: list[str] = field(default_factory=list)
    points: Any = None
    correct_answer: Any = ""
    explanation: str = ""
    grading_guidance: str = ""
    deduction_points: str = ""


@dataclass
class StudentSubmissionExport:
    content: bytes
    filename: str


def build_student_submission_export_docx(
    conn,
    *,
    assignment_id: str,
    student_pk_id: int,
) -> StudentSubmissionExport:
    context = _load_export_context(conn, assignment_id=assignment_id, student_pk_id=student_pk_id)
    submission = context["submission"]
    if str(submission.get("submission_status") or "").strip().lower() != "graded":
        raise HTTPException(400, "批改完成后才能导出复习 Word")
    if int(submission.get("is_absence_score") or 0):
        raise HTTPException(400, "缺交记 0 的记录没有可导出的学生答卷")

    submission_files = _load_submission_files(conn, int(submission["submission_id"]))
    answer_items = _normalize_answer_items(submission.get("answers_json"))
    pages = _build_export_pages(context, answer_items)
    feedback_by_question = _extract_question_feedback(submission.get("feedback_md"))

    document = Document()
    _setup_document(document, context)
    _write_title_block(document, context)
    _write_overview(document, context, pages)
    _write_questions(document, pages, answer_items, submission_files, feedback_by_question)
    _write_unmatched_attachments(document, submission_files)

    buffer = io.BytesIO()
    document.save(buffer)
    student_name = str(context["student"].get("student_name") or "学生").strip()
    student_number = str(context["student"].get("student_id_number") or "无学号").strip()
    assignment_title = str(context["assignment"].get("assignment_title") or "作业考试").strip()
    filename = _safe_filename(f"{assignment_title}_{student_name}_{student_number}_复习导出.docx")
    return StudentSubmissionExport(content=buffer.getvalue(), filename=filename)


def _load_export_context(conn, *, assignment_id: str, student_pk_id: int) -> dict[str, dict[str, Any]]:
    row = conn.execute(
        """
        SELECT
            s.id AS submission_id,
            s.status AS submission_status,
            s.score,
            s.feedback_md,
            s.answers_json,
            s.submitted_at,
            s.started_at,
            s.is_absence_score,
            s.score_before_late_penalty,
            s.late_penalty_points,
            s.late_score_cap_applied,
            a.id AS assignment_id,
            a.title AS assignment_title,
            a.requirements_md,
            a.rubric_md,
            a.exam_paper_id,
            a.course_id,
            a.class_offering_id,
            c.name AS course_name,
            ep.title AS paper_title,
            ep.description AS paper_description,
            ep.questions_json,
            ep.exam_config_json,
            stu.name AS student_name,
            stu.student_id_number
        FROM submissions s
        JOIN assignments a ON a.id = s.assignment_id
        JOIN students stu ON stu.id = s.student_pk_id
        LEFT JOIN courses c ON c.id = a.course_id
        LEFT JOIN exam_papers ep ON ep.id = a.exam_paper_id
        WHERE s.assignment_id = ?
          AND s.student_pk_id = ?
        LIMIT 1
        """,
        (str(assignment_id), int(student_pk_id)),
    ).fetchone()
    if not row:
        raise HTTPException(404, "未找到该学生的提交记录")

    item = dict(row)
    return {
        "submission": item,
        "assignment": item,
        "student": item,
        "paper": item,
    }


def _load_submission_files(conn, submission_id: int) -> list[ExportAttachment]:
    rows = conn.execute(
        """
        SELECT id, original_filename, relative_path, stored_path, mime_type, file_size, file_ext
        FROM submission_files
        WHERE submission_id = ?
        ORDER BY COALESCE(relative_path, original_filename), id
        """,
        (int(submission_id),),
    ).fetchall()
    attachments: list[ExportAttachment] = []
    for row in rows:
        item = dict(row)
        relative_path = str(item.get("relative_path") or item.get("original_filename") or "").replace("\\", "/")
        name = str(item.get("original_filename") or Path(relative_path).name or "附件").strip()
        mime_type = str(item.get("mime_type") or "").lower()
        file_ext = str(item.get("file_ext") or Path(name).suffix or Path(relative_path).suffix or "").lower()
        resolved = resolve_submission_file_path(str(item.get("stored_path") or ""))
        path = Path(resolved) if resolved else None
        is_image = bool(mime_type.startswith("image/") or file_ext in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"})
        attachments.append(
            ExportAttachment(
                name=name,
                relative_path=relative_path,
                mime_type=mime_type,
                file_ext=file_ext,
                path=path,
                is_image=is_image,
            )
        )
    return attachments


def _load_json_object(raw: Any) -> dict[str, Any]:
    if isinstance(raw, dict):
        return raw
    if not raw:
        return {}
    try:
        parsed = json.loads(str(raw))
    except (TypeError, json.JSONDecodeError):
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _build_export_pages(context: dict[str, dict[str, Any]], answer_items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    assignment = context["assignment"]
    paper_data = _load_json_object(context["paper"].get("questions_json"))
    raw_pages: Any = None
    if isinstance(paper_data.get("questions"), dict) and isinstance(paper_data["questions"].get("pages"), list):
        raw_pages = paper_data["questions"]["pages"]
    elif isinstance(paper_data.get("pages"), list):
        raw_pages = paper_data["pages"]
    elif isinstance(paper_data.get("questions"), list):
        raw_pages = [{"name": "试卷题目", "questions": paper_data["questions"]}]

    pages: list[dict[str, Any]] = []
    question_number = 1
    if isinstance(raw_pages, list):
        for page_index, page in enumerate(raw_pages, start=1):
            if not isinstance(page, dict):
                continue
            questions = []
            for raw_question in page.get("questions") or []:
                if not isinstance(raw_question, dict):
                    continue
                question = _export_question_from_exam_item(
                    raw_question,
                    section_name=str(page.get("name") or page.get("title") or f"第{page_index}部分"),
                    number=question_number,
                )
                questions.append(question)
                question_number += 1
            if questions:
                pages.append({"name": str(page.get("name") or page.get("title") or f"第{page_index}部分"), "questions": questions})

    if pages:
        return pages

    questions = []
    if answer_items:
        for item in answer_items:
            label = str(item.get("question") or item.get("question_id") or item.get("title") or f"第 {question_number} 题").strip()
            questions.append(
                ExportQuestion(
                    section_name="作业答题",
                    number=question_number,
                    question_id=str(item.get("question_id") or question_number),
                    question_type=str(item.get("type") or "assignment"),
                    text=label,
                    correct_answer=assignment.get("rubric_md") or "教师未配置独立标准答案，请结合评分标准与批改评语复习。",
                )
            )
            question_number += 1
    else:
        questions.append(
            ExportQuestion(
                section_name="作业答题",
                number=1,
                question_id="assignment",
                question_type="assignment",
                text=assignment.get("requirements_md") or assignment.get("assignment_title") or "作业要求",
                correct_answer=assignment.get("rubric_md") or "教师未配置独立标准答案，请结合评分标准与批改评语复习。",
            )
        )
    return [{"name": "作业答题", "questions": questions}]


def _export_question_from_exam_item(raw: dict[str, Any], *, section_name: str, number: int) -> ExportQuestion:
    question_id = str(raw.get("id") or raw.get("question_id") or number).strip()
    correct = (
        raw.get("answer")
        if "answer" in raw
        else raw.get("correct_answer", raw.get("reference_answer", raw.get("standard_answer", "")))
    )
    return ExportQuestion(
        section_name=section_name,
        number=number,
        question_id=question_id,
        question_type=str(raw.get("type") or raw.get("question_type") or ""),
        text=str(raw.get("text") or raw.get("question") or raw.get("title") or raw.get("stem") or "").strip(),
        options=[str(option) for option in (raw.get("options") or raw.get("choices") or []) if str(option).strip()],
        points=raw.get("points"),
        correct_answer=correct,
        explanation=str(raw.get("explanation") or raw.get("analysis") or "").strip(),
        grading_guidance=str(raw.get("grading_guidance") or raw.get("guidance") or raw.get("score_points") or "").strip(),
        deduction_points=str(raw.get("deduction_points") or raw.get("deductions") or "").strip(),
    )


def _normalize_answer_items(raw_answers_json: Any) -> list[dict[str, Any]]:
    payload: Any = raw_answers_json
    if isinstance(raw_answers_json, str):
        try:
            payload = json.loads(raw_answers_json) if raw_answers_json.strip() else {}
        except json.JSONDecodeError:
            return [{"question_id": "raw", "question": "完整答案", "answer": raw_answers_json}]
    answers = payload.get("answers", payload) if isinstance(payload, dict) else payload
    if isinstance(answers, list):
        return [dict(item) for item in answers if isinstance(item, dict)]
    if isinstance(answers, dict):
        items: list[dict[str, Any]] = []
        for key, value in answers.items():
            if isinstance(value, dict):
                item = dict(value)
                item.setdefault("question_id", key)
            else:
                item = {"question_id": key, "question": key, "answer": value}
            items.append(item)
        return items
    return []


def _setup_document(document: Document, context: dict[str, dict[str, Any]]) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(0.85)

    styles = document.styles
    for style_name in ("Normal", "Body Text"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = BODY_FONT
            style.font.size = Pt(10.5)
            _set_style_east_asia_font(style, BODY_FONT)
    for style_name, size, color in (
        ("Title", 18, "1F2937"),
        ("Heading 1", 14, "1F2937"),
        ("Heading 2", 12, "1F2937"),
        ("Heading 3", 11, "374151"),
    ):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = BODY_FONT
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            _set_style_east_asia_font(style, BODY_FONT)

    student = context["student"]
    watermark = f"{student.get('student_name') or '学生'}  {student.get('student_id_number') or '无学号'}"
    _add_watermark(section, watermark)
    _add_page_footer(section)


def _write_title_block(document: Document, context: dict[str, dict[str, Any]]) -> None:
    assignment = context["assignment"]
    student = context["student"]
    submission = context["submission"]
    title = str(assignment.get("paper_title") or assignment.get("assignment_title") or "作业考试复习导出").strip()

    title_para = document.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    _set_run_font(title_run, size=18, bold=True, color="111827")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs_with_inline_formatting(
        subtitle,
        f"{student.get('student_name') or '学生'}（{student.get('student_id_number') or '无学号'}）复习打印版",
        size=10,
        color="6B7280",
    )
    subtitle.paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table, fill="F8FAFC")
    meta_rows = [
        ("课程", assignment.get("course_name") or "-", "作业/考试", assignment.get("assignment_title") or title),
        ("学生", student.get("student_name") or "-", "学号", student.get("student_id_number") or "-"),
        ("提交时间", _format_datetime(submission.get("submitted_at")), "导出时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("成绩", _score_label(submission), "用途", "打印复习"),
    ]
    for row_values in meta_rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value or "-"))
            if index in {0, 2}:
                _set_run_font(run, size=9, bold=True, color="475569")
                _shade_cell(cells[index], "EEF2F7")
            else:
                _set_run_font(run, size=9, color="111827")
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _write_overview(document: Document, context: dict[str, dict[str, Any]], pages: list[dict[str, Any]]) -> None:
    assignment = context["assignment"]
    submission = context["submission"]
    document.add_paragraph()
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(4)
    heading.add_run("复习导读")
    _add_markdown_blocks(
        document,
        "本文档按原题顺序整理题目、学生作答、参考答案与必要附件，便于打印后离线复习。题目和参考答案是主要内容，批改反馈只作为辅助订正线索。",
        font_size=10.5,
    )

    summary_lines = []
    if assignment.get("paper_description"):
        summary_lines.append(str(assignment.get("paper_description") or "").strip())
    if assignment.get("requirements_md") and not assignment.get("exam_paper_id"):
        summary_lines.append("作业要求：\n" + str(assignment.get("requirements_md") or "").strip())
    if submission.get("feedback_md"):
        overview = _extract_feedback_overview(str(submission.get("feedback_md") or ""))
        if overview:
            summary_lines.append("批改总览：\n" + overview)
    if summary_lines:
        _add_markdown_blocks(document, "\n\n".join(summary_lines), font_size=10)

    total_questions = sum(len(page.get("questions") or []) for page in pages)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run(f"共 {len(pages)} 个部分，{total_questions} 道题。")
    _set_run_font(run, size=9, color="64748B")


def _write_questions(
    document: Document,
    pages: list[dict[str, Any]],
    answer_items: list[dict[str, Any]],
    submission_files: list[ExportAttachment],
    feedback_by_question: dict[str, str],
) -> None:
    for page in pages:
        section_heading = document.add_paragraph(style="Heading 1")
        section_heading.paragraph_format.space_before = Pt(10)
        section_heading.add_run(str(page.get("name") or "试题"))
        for question in page.get("questions") or []:
            _write_question(document, question, answer_items, submission_files, feedback_by_question)


def _write_question(
    document: Document,
    question: ExportQuestion,
    answer_items: list[dict[str, Any]],
    submission_files: list[ExportAttachment],
    feedback_by_question: dict[str, str],
) -> None:
    answer_item = _find_answer_for_question(question, answer_items)
    student_answer = _answer_text(answer_item)
    attachments = _attachments_for_answer(answer_item, submission_files)
    correct_answer = _correct_answer_text(question)

    heading = document.add_paragraph(style="Heading 2")
    heading.paragraph_format.space_before = Pt(9)
    heading.paragraph_format.space_after = Pt(4)
    label = f"{question.number}. {question.question_id}"
    if question.points not in (None, ""):
        label += f"（{question.points} 分）"
    heading_run = heading.add_run(label)
    _set_run_font(heading_run, size=12, bold=True, color="111827")

    _add_markdown_blocks(document, question.text or "(题干为空)", font_size=10.5)
    if question.options:
        for option in question.options:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs_with_inline_formatting(paragraph, option, size=10)

    _add_label(document, "学生回答", "2563EB")
    if student_answer:
        _add_markdown_blocks(document, student_answer, font_size=10)
    else:
        _add_empty_text(document, "(未作答)")
    if attachments:
        _write_question_attachments(document, attachments)

    _add_label(document, "参考答案", "047857")
    _add_markdown_blocks(document, correct_answer or "暂未配置标准答案，请结合教师讲解与批改评语复习。", font_size=10.5)

    guidance_parts = []
    if question.explanation:
        guidance_parts.append("解析：" + question.explanation)
    if question.grading_guidance:
        guidance_parts.append("评分要点：" + question.grading_guidance)
    if question.deduction_points:
        guidance_parts.append("易扣分点：" + question.deduction_points)
    if guidance_parts:
        _add_label(document, "订正提示", "64748B")
        _add_markdown_blocks(document, "\n\n".join(guidance_parts), font_size=9.5, color="475569")

    feedback = feedback_by_question.get(_feedback_key(question.question_id)) or feedback_by_question.get(str(question.number))
    if feedback:
        _add_label(document, "本题批改反馈", "B45309")
        _add_markdown_blocks(document, feedback, font_size=9.5, color="6B4E16")

    _add_divider(document)


def _write_question_attachments(document: Document, attachments: list[ExportAttachment]) -> None:
    image_index = 1
    listed_non_images: list[str] = []
    for attachment in attachments:
        if attachment.is_image and attachment.path and attachment.path.exists():
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.keep_with_next = True
            caption_run = caption.add_run(f"学生附图 {image_index}: {attachment.name}")
            _set_run_font(caption_run, size=8.5, color="64748B")
            try:
                width_in = _fit_image_width(attachment.path)
                document.add_picture(str(attachment.path), width=Inches(width_in))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_index += 1
            except Exception:
                listed_non_images.append(f"{attachment.name}（图片插入失败，可在平台原附件中查看）")
        else:
            listed_non_images.append(attachment.name)
    if listed_non_images:
        paragraph = document.add_paragraph()
        _add_runs_with_inline_formatting(paragraph, "本题附件：" + "；".join(listed_non_images), size=9, color="64748B")


def _write_unmatched_attachments(document: Document, submission_files: list[ExportAttachment]) -> None:
    unmatched = [item for item in submission_files if not item.matched_to_question]
    if not unmatched:
        return
    document.add_paragraph(style="Heading 1").add_run("其他提交附件")
    for item in unmatched:
        paragraph = document.add_paragraph(style="List Bullet")
        _add_runs_with_inline_formatting(paragraph, item.relative_path or item.name, size=9.5, color="475569")


def _find_answer_for_question(question: ExportQuestion, answer_items: list[dict[str, Any]]) -> dict[str, Any] | None:
    question_id_key = _feedback_key(question.question_id)
    for index, item in enumerate(answer_items, start=1):
        candidates = [
            item.get("question_id"),
            item.get("id"),
            item.get("question_no"),
            item.get("question"),
            index,
        ]
        for candidate in candidates:
            if _feedback_key(candidate) == question_id_key:
                return item
        if str(item.get("question") or "").strip() == question.text.strip():
            return item
        if index == question.number:
            fallback = item
    return fallback if "fallback" in locals() else None


def _answer_text(item: dict[str, Any] | None) -> str:
    if not item:
        return ""
    for key in ("answer", "content", "text", "value"):
        if key in item:
            return _value_to_markdown(item.get(key))
    without_attachments = {k: v for k, v in item.items() if k not in {"attachments", "question", "question_id", "type"}}
    return _value_to_markdown(without_attachments) if without_attachments else ""


def _attachments_for_answer(item: dict[str, Any] | None, submission_files: list[ExportAttachment]) -> list[ExportAttachment]:
    if not item or not isinstance(item.get("attachments"), list):
        return []
    matched: list[ExportAttachment] = []
    seen: set[str] = set()
    for raw_attachment in item.get("attachments") or []:
        if not isinstance(raw_attachment, dict):
            continue
        attachment = _match_submission_file(raw_attachment, submission_files)
        if not attachment:
            continue
        identity = attachment.relative_path.lower() or attachment.name.lower()
        if identity in seen:
            continue
        seen.add(identity)
        attachment.matched_to_question = True
        matched.append(attachment)
    return matched


def _match_submission_file(raw_attachment: dict[str, Any], submission_files: list[ExportAttachment]) -> ExportAttachment | None:
    keys = [
        raw_attachment.get("relative_path"),
        raw_attachment.get("stored_relative_path"),
        raw_attachment.get("file_name"),
        raw_attachment.get("filename"),
    ]
    normalized_keys = [str(key or "").replace("\\", "/").strip().lower() for key in keys if str(key or "").strip()]
    for file in submission_files:
        candidates = [
            file.relative_path,
            file.name,
            Path(file.relative_path).name if file.relative_path else "",
        ]
        normalized_candidates = [str(candidate or "").replace("\\", "/").strip().lower() for candidate in candidates if str(candidate or "").strip()]
        if any(candidate == key or candidate.endswith("/" + key) for candidate in normalized_candidates for key in normalized_keys):
            return file
    return None


def _correct_answer_text(question: ExportQuestion) -> str:
    answer = _value_to_markdown(question.correct_answer)
    if answer:
        return answer
    if question.grading_guidance:
        return question.grading_guidance
    if question.explanation:
        return question.explanation
    return ""


def _value_to_markdown(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(f"- {_value_to_markdown(item)}" for item in value if _value_to_markdown(item).strip())
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value or "").strip()


def _add_markdown_blocks(
    container: Any,
    text: str,
    *,
    font_size: float = 10.5,
    color: str = "111827",
) -> None:
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        _add_runs_with_inline_formatting(paragraph, " ".join(item.strip() for item in paragraph_buffer).strip(), size=font_size, color=color)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            _add_code_block(container, "\n".join(code_lines), font_size=font_size - 1)
            continue

        if _is_markdown_table_start(lines, index):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and "|" in lines[index]:
                table_lines.append(lines[index])
                index += 1
            _add_markdown_table(container, table_lines, font_size=font_size - 1)
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            paragraph = container.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4 if level <= 2 else 2)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_runs_with_inline_formatting(paragraph, heading_match.group(2), size=max(font_size, 10.5) + max(0, 3 - level), bold=True, color=color)
            index += 1
            continue

        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        ordered_match = re.match(r"^\d+[\.)]\s+(.+)$", stripped)
        if bullet_match or ordered_match:
            flush_paragraph()
            paragraph = container.add_paragraph(style="List Number" if ordered_match else "List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            _add_runs_with_inline_formatting(paragraph, (ordered_match or bullet_match).group(1), size=font_size, color=color)
            index += 1
            continue

        quote_match = re.match(r"^>\s*(.+)$", stripped)
        if quote_match:
            flush_paragraph()
            paragraph = container.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.45)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_runs_with_inline_formatting(paragraph, quote_match.group(1), size=font_size, italic=True, color="64748B")
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def _add_runs_with_inline_formatting(
    paragraph,
    text: str,
    *,
    size: float = 10.5,
    color: str = "111827",
    bold: bool = False,
    italic: bool = False,
) -> None:
    if not text:
        return
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            _set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size - 0.5, color="7C2D12", bold=bold, italic=italic, font_name=MONO_FONT)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, color=color, bold=True, italic=italic)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, color=color, bold=bold, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def _add_code_block(container: Any, code: str, *, font_size: float = 9) -> None:
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(code or " ")
    _set_run_font(run, size=font_size, color="334155", font_name=MONO_FONT)
    _shade_paragraph(paragraph, "F8FAFC")


def _is_markdown_table_start(lines: Sequence[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    next_line = lines[index + 1].strip()
    return "|" in current and bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", next_line))


def _parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _add_markdown_table(container: Any, table_lines: list[str], *, font_size: float = 9) -> None:
    if len(table_lines) < 2:
        return
    header = _parse_table_row(table_lines[0])
    rows = [_parse_table_row(row) for row in table_lines[2:]]
    column_count = max(len(header), *(len(row) for row in rows)) if rows else len(header)
    table = container.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)
    for col_index in range(column_count):
        cell = table.rows[0].cells[col_index]
        _shade_cell(cell, "EAF2FF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        _add_runs_with_inline_formatting(paragraph, header[col_index] if col_index < len(header) else "", size=font_size, bold=True)
    _mark_header_row_repeat(table)
    for row in rows:
        cells = table.add_row().cells
        for col_index in range(column_count):
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[col_index].paragraphs[0]
            _add_runs_with_inline_formatting(paragraph, row[col_index] if col_index < len(row) else "", size=font_size)
    container.add_paragraph()


def _add_label(document: Document, label: str, color: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    _set_run_font(run, size=9.5, bold=True, color=color)


def _add_empty_text(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, size=10, italic=True, color="94A3B8")


def _add_divider(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    font_name: str = BODY_FONT,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT if font_name == BODY_FONT else font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_east_asia_font(style, font_name: str) -> None:
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_table_borders(table, *, fill: str | None = None) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CBD5E1")
        borders.append(border)
    tbl_pr.append(borders)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tbl_pr.append(shading)


def _mark_header_row_repeat(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_watermark(section, text: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    watermark_bytes = _build_watermark_png(text)
    if watermark_bytes:
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(watermark_bytes), width=Cm(17.2))
        return
    run = paragraph.add_run(text)
    _set_run_font(run, size=16, color="D1D5DB")


def _build_watermark_png(text: str) -> bytes | None:
    try:
        canvas_width, canvas_height = 1600, 2200
        canvas = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
        text_layer = Image.new("RGBA", (canvas_width, 420), (255, 255, 255, 0))
        draw = ImageDraw.Draw(text_layer)
        font = _load_watermark_font(72)
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        draw.text(
            ((canvas_width - text_width) / 2, (420 - text_height) / 2),
            text,
            font=font,
            fill=(100, 116, 139, 26),
        )
        rotated = text_layer.rotate(315, expand=True, resample=Image.Resampling.BICUBIC)
        canvas.alpha_composite(
            rotated,
            ((canvas_width - rotated.width) // 2, (canvas_height - rotated.height) // 2),
        )
        output = io.BytesIO()
        canvas.save(output, format="PNG")
        return output.getvalue()
    except Exception:
        return None


def _load_watermark_font(size: int):
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size)
            except Exception:
                continue
    return ImageFont.load_default()


def _add_page_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(paragraph.add_run("第 "), size=9, color="64748B")
    _add_field(paragraph, "PAGE", fallback="1")
    _set_run_font(paragraph.add_run(" 页"), size=9, color="64748B")


def _add_field(paragraph, instruction: str, *, fallback: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {instruction} \\* MERGEFORMAT "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(fallback)
    _set_run_font(result_run, size=9, color="64748B")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _fit_image_width(path: Path) -> float:
    try:
        with Image.open(path) as image:
            width_px, height_px = image.size
    except Exception:
        return QUESTION_MAX_IMAGE_WIDTH_IN
    if width_px <= 0 or height_px <= 0:
        return QUESTION_MAX_IMAGE_WIDTH_IN
    aspect = width_px / height_px
    width_in = min(QUESTION_MAX_IMAGE_WIDTH_IN, QUESTION_MAX_IMAGE_HEIGHT_IN * aspect)
    return max(1.5, width_in)


def _extract_feedback_overview(feedback_md: str) -> str:
    text = str(feedback_md or "").strip()
    if not text:
        return ""
    match = re.search(r"^\s*##\s*逐题反馈\s*$", text, flags=re.MULTILINE)
    overview = text[: match.start()].strip() if match else text
    overview = re.sub(r"^\s*##\s*总览评语\s*", "", overview, flags=re.MULTILINE).strip()
    return overview[:900]


def _extract_question_feedback(feedback_md: Any) -> dict[str, str]:
    text = str(feedback_md or "").replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        return {}
    matches = list(re.finditer(r"^\s*###\s*(?:第\s*)?(.+?)(?:\s*题)?\s*$", text, flags=re.MULTILINE))
    result: dict[str, str] = {}
    for index, match in enumerate(matches):
        key = _feedback_key(match.group(1))
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if key and body:
            result[key] = body
            result[str(index + 1)] = body
    return result


def _feedback_key(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"^(第\s*)", "", text)
    text = re.sub(r"(题|question|q)$", "", text).strip()
    return re.sub(r"\s+", "", text)


def _format_datetime(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return "-"
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M"):
        try:
            return datetime.fromisoformat(text).strftime("%Y-%m-%d %H:%M")
        except ValueError:
            continue
    return text[:16]


def _score_label(submission: dict[str, Any]) -> str:
    score = submission.get("score")
    if score in (None, ""):
        return "已批改"
    label = f"{score} 分"
    if submission.get("score_before_late_penalty") not in (None, ""):
        label += f"（原始 {submission.get('score_before_late_penalty')}，补交扣 {submission.get('late_penalty_points') or 0}）"
    return label


def _safe_filename(name: str) -> str:
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(name or "export.docx"))
    safe = re.sub(r"_+", "_", safe).strip("._ ")
    return safe or "export.docx"
